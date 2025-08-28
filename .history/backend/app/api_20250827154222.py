from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Depends, Query
from fastapi.responses import FileResponse, JSONResponse
import os, time
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
from sqlalchemy.orm import Session, joinedload
from sqlalchemy.future import select
from sqlalchemy.ext.asyncio import AsyncSession
from pydantic import ValidationError
from typing import List, Optional
from datetime import datetime
import pytz
import traceback
import urllib

# 导入自定义模块
from .utils import save_uploaded_file
from .AI_client import generate_text_for_user, get_user_model_preference
from .database import SessionLocal, get_db, get_async_db 
from .models import (
    DocumentHistory, Template, AIModel, Conversation, Message, Platform,
    AIModelResponse, PlatformModelResponse, User, AIModelCreate, SystemModelResponse, Model, Notification
)
from .deps import get_current_user
from .encryption import encrypt_api_key, decrypt_api_key

# 加载环境变量
load_dotenv()

router = APIRouter()

# ----------------- 目录配置 -----------------
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), '..', 'uploads')
os.makedirs(UPLOAD_DIR, exist_ok=True)

DOWNLOAD_DIR = os.path.join(os.getcwd(), "generated_docs")
os.makedirs(DOWNLOAD_DIR, exist_ok=True)

# ----------------- 公文生成Prompt配置 -----------------
PROMPTS = {
    "通知": "你是专业的正式公文写作助手，需严格按照《党政机关公文处理工作条例》规定的通知格式撰写。内容需包含标题（发文机关+事由+文种）、主送机关、正文（明确通知缘由、具体事项、执行要求，涉及任务需标注责任主体）、附件说明（如有附件需注明附件名称及数量）、发文机关署名、成文日期、签发人（上行文需标注），语言需庄重严谨、逻辑清晰，避免歧义。",
    "请示": "你是专业的正式公文写作助手，需严格遵循请示“一文一事、单头主送”的核心要求撰写。内容需包含标题（发文机关+事由+文种）、主送机关（仅限一个直接上级机关）、正文（开门见山说明请示缘由，需引用政策依据或实际需求；清晰列出请示事项，做到具体明确；提出合理可行的建议方案供上级决策）、发文机关署名、成文日期、签发人，语言需谦恭得体，避免多头请示或夹带报告事项。",
    "会议纪要": "你是专业的正式公文写作助手，需按规范会议纪要格式撰写，确保内容客观准确、要素完整。需包含标题（会议名称+文种，如“XX会议纪要”）、基本信息（会议时间、会议地点、主持人、记录人、参会人员/列席人员/缺席人员名单及职务）、正文（分点梳理会议讨论要点，需准确反映不同观点；明确会议形成的决议事项，标注每项决议的责任单位、完成时限）、结尾（如需补充说明未尽事宜或后续沟通方式，可简要标注），语言需简洁凝练，避免主观评价。",
    "报告": "你是专业的正式公文写作助手，需严格按照报告的法定格式与写作要求撰写，区分“汇报工作”“反映情况”“答复询问”等不同报告类型。内容需包含标题（发文机关+事由+文种，如“XX关于XX工作的报告”）、主送机关、正文（开头简要说明报告背景或依据，主体部分详细阐述工作进展、成果、存在问题、原因分析及下一步计划，答复类报告需直接回应上级询问事项）、发文机关署名、成文日期、签发人（上行文需标注），语言需客观真实，不得夹带请示事项。",
    "函": "你是专业的正式公文写作助手，需根据函的“不相隶属机关商洽工作、询问答复问题、请求批准事项”功能，按规范格式撰写。内容需包含标题（发文机关+事由+文种，如“XX关于XX事宜的函”或“XX关于XX问题的复函”）、主送机关（明确的不相隶属机关）、正文（商洽类需说明事由及具体需求，询问类需列出清晰的问题清单，答复类需直接回应来函事项，请求批准类需说明依据及具体请求）、发文机关署名、成文日期，语言需平和得体，避免命令性或指令性表述。",
    "对话": "你是专业的公文场景对话撰写助手，需围绕公文办理全流程（如公文起草沟通、审批意见反馈、事项协调对接等）撰写对话内容。需明确对话场景（如“公文起草小组沟通会议对话”“上下级机关审批意见反馈对话”）、对话主体（标注角色及职务，如“起草人-XX部科员”“审批人-XX局副局长”）、对话逻辑（需符合公文办理规范，内容需聚焦具体事项，如格式修改建议、内容补充要求、办理时限确认等），语言需贴合职场沟通语境，既保持正式性，又体现沟通的针对性与高效性。"
}

# ----------------- 核心接口：上传模板 -----------------
@router.post("/upload-template")
async def upload_template(
    file: UploadFile = File(...),
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # 保存模板文件到本地目录
    file_path = await save_uploaded_file(file, UPLOAD_DIR)
    filename = os.path.basename(file_path)
    
    # 读取模板内容（存入数据库）
    template_content = ""
    try:
        doc = Document(file_path)
        template_content = "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"读取模板内容失败：{str(e)}")
    
    # 保存模板记录到数据库
    template = Template(
        user_id=current_user.id,
        filename=filename,
        original_name=file.filename,
        content=template_content,
        status="active"
    )
    db.add(template)
    db.commit()
    db.refresh(template)
    
    return {
        "id": template.id,
        "filename": template.filename,
        "original_name": template.original_name,
        "uploaded_at": template.uploaded_at
    }

# ----------------- 核心接口：公文生成 -----------------
@router.post("/generate")
async def generate_document(
    doc_type: str = Form(...),
    user_input: str = Form(...),
    conv_id: Optional[int] = Form(None),
    ai_model_id: Optional[int] = Form(None),
    template_id: Optional[int] = Form(None),
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # 1. 组装生成Prompt
    base_prompt = PROMPTS.get(doc_type, f"请写一份正式公文：{doc_type}")
    if template_id:
        template = db.query(Template).filter(
            Template.id == template_id,
            Template.user_id == current_user.id,
            Template.status == "active"
        ).first()
        if not template:
            raise HTTPException(status_code=404, detail="指定模板不存在或无权访问")
        prompt = f"{base_prompt}\n模板内容：{template.content}\n用户要求：{user_input}"
    else:
        prompt = f"{base_prompt}\n用户要求：{user_input}"

    # 2. 获取用户AI模型配置
    user_config = None
    selected_ai_model = None
    
    if ai_model_id:
        selected_ai_model = db.query(AIModel).filter(
            AIModel.id == ai_model_id,
            AIModel.user_id == current_user.id
        ).options(joinedload(AIModel.model)).first()
        if not selected_ai_model:
            raise HTTPException(status_code=404, detail="选择的AI模型不存在或无权访问")
        
        # 平台名称映射
        platform_mapping = {
            "Alibaba": "qwen",
            "OpenAI": "openai",
            "Anthropic": "anthropic",
            "Google": "gemini"
        }
        db_platform_name = selected_ai_model.model.platform.name
        provider = platform_mapping.get(db_platform_name, db_platform_name.lower())
        
        user_config = {
            "provider": provider,
            "api_key": decrypt_api_key(selected_ai_model.api_key),
            "model_name": selected_ai_model.model.name,
            "base_url": selected_ai_model.effective_base_url
        }
    else:
        user_config = get_user_model_preference(user_id=current_user.id, db=db)
        if not user_config:
            raise HTTPException(status_code=400, detail="请先在「API Keys管理」中添加AI模型配置")
        
        platform_mapping = {
            "Alibaba": "qwen",
            "OpenAI": "openai",
            "Anthropic": "anthropic",
            "Google": "gemini"
        }
        user_config["provider"] = platform_mapping.get(
            user_config["provider"], 
            user_config["provider"].lower()
        )
        
        selected_ai_model = db.query(AIModel).filter(
            AIModel.model.has(
                Model.platform.has(name=user_config["provider"]),
                name=user_config["model_name"]
            ),
            AIModel.user_id == current_user.id
        ).options(joinedload(AIModel.model)).first()
        
        if not selected_ai_model:
            raise HTTPException(status_code=500, detail="用户默认模型配置异常，请重新添加")

    # 3. 调用多平台Client生成公文内容
    try:
        generated = generate_text_for_user(
            user_id=current_user.id,
            prompt=prompt,
            db=db,
            system_prompt=base_prompt,
            user_config=user_config
        )
    except Exception as e:
        error_msg = str(e)
        if "401" in error_msg and "invalid_api_key" in error_msg.lower():
            raise HTTPException(
                status_code=401, 
                detail="API Key不正确或已失效，请检查并重新配置"
            )
        elif "不支持的平台" in error_msg:
            raise HTTPException(
                status_code=400,
                detail="所选AI平台暂不支持，请更换其他平台"
            )
        else:
            raise HTTPException(
                status_code=500, 
                detail=f"AI生成公文失败：{error_msg}"
            )

    # 4. 生成DOCX文件
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'FangSong_GB2312'
    font.size = Pt(16)
    font.bold = False
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong_GB2312')
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    for line in generated.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        else:
            p = doc.add_paragraph()
            parts = line.split("**")
            for i, part in enumerate(parts):
                if i % 2 == 1:
                    run = p.add_run(part)
                    run.bold = True
                else:
                    subparts = part.split("*")
                    for j, sp in enumerate(subparts):
                        run = p.add_run(sp)
                        if j % 2 == 1:
                            run.italic = True

    # 保存DOCX文件
    filename = f"{doc_type}_{current_user.id}_{int(time.time())}.docx"
    file_path = os.path.join(DOWNLOAD_DIR, filename)
    doc.save(file_path)

    # 5. 保存公文记录
    doc_record = DocumentHistory(
        user_id=current_user.id,
        doc_type=doc_type,
        content=generated,
        filename=filename,
        template_id=template_id
    )
    db.add(doc_record)
    db.commit()
    db.refresh(doc_record)

    # 6. 关联数据库会话
    if conv_id:
        conversation = db.query(Conversation).filter(
            Conversation.id == conv_id,
            Conversation.user_id == current_user.id
        ).first()
        if not conversation:
            raise HTTPException(status_code=404, detail="指定会话不存在或无权访问")
        conversation.updated_at = datetime.now(pytz.UTC)
        conversation.ai_model_id = selected_ai_model.id
        conversation.status = "active"
    else:
        title = f"{doc_type}生成_{user_input.strip()[:8]}..." if len(user_input.strip()) > 8 else f"{doc_type}生成"
        conversation = Conversation(
            user_id=current_user.id,
            ai_model_id=selected_ai_model.id,
            title=title,
            status="active",
            created_at=datetime.now(pytz.UTC),
            updated_at=datetime.now(pytz.UTC)
        )
        db.add(conversation)
    db.commit()
    db.refresh(conversation)

    # 7. 保存消息
    user_message = Message(
        conversation_id=conversation.id,
        role="user",
        content=user_input,
        created_at=datetime.now(pytz.UTC)
    )
    ai_message = Message(
        conversation_id=conversation.id,
        role="assistant",
        content=generated,
        docx_file=filename,
        created_at=datetime.now(pytz.UTC)
    )
    db.add_all([user_message, ai_message])
    db.commit()

    # 8. 返回结果
    return JSONResponse({
        "text": generated,
        "filename": filename,
        "conv_id": conversation.id,
        "doc_id": doc_record.id,
        "used_model": f"{selected_ai_model.model.platform.name} - {selected_ai_model.model.name}"
    }, status_code=200)

# ----------------- 接口：下载生成的DOCX文件 -----------------
@router.get("/download/{filename}")
async def download(
    filename: str, 
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    decoded_filename = urllib.parse.unquote(filename)
    file_path = os.path.join(DOWNLOAD_DIR, decoded_filename)
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件未找到或已被删除")
    
    doc_record = db.query(DocumentHistory).filter(
        DocumentHistory.filename == decoded_filename,
        DocumentHistory.user_id == current_user.id
    ).first()
    if not doc_record:
        raise HTTPException(status_code=403, detail="无权访问此文件（非本人生成）")
    
    encoded_filename = urllib.parse.quote(decoded_filename)
    
    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
        }
    )



# ----------------- 接口：获取用户模板列表 -----------------
@router.get("/templates")
async def get_templates(
    current_user = Depends(get_current_user), 
    db: Session = Depends(get_db)
):
    templates = db.query(Template).filter(
        Template.user_id == current_user.id,
        Template.status == "active"
    ).order_by(Template.uploaded_at.desc()).all()
    
    return [
        {
            "id": t.id,
            "filename": t.filename,
            "original_name": t.original_name,
            "uploaded_at": t.uploaded_at.strftime("%Y-%m-%d %H:%M:%S"),
            "content_preview": t.content[:30] + "..." if len(t.content) > 30 else t.content
        }
        for t in templates
    ]

# ----------------- 接口：获取模板内容 -----------------
@router.get("/template-content/{template_id}")
async def get_template_content(
    template_id: int, 
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    template = db.query(Template).filter(
        Template.id == template_id,
        Template.user_id == current_user.id,
        Template.status == "active"
    ).first()
    
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在或无权访问")
    
    return JSONResponse({
        "id": template.id,
        "original_name": template.original_name,
        "content": template.content
    })

# ----------------- 接口：获取用户API Key列表 -----------------
@router.get("/keys", response_model=List[AIModelResponse])
async def get_user_ai_configs(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_async_db)
):
    try:
        stmt = select(AIModel).filter(
            AIModel.user_id == current_user.id
        ).options(
            joinedload(AIModel.model).joinedload(Model.platform)
        ).order_by(AIModel.created_at.desc())
        
        result = await db.execute(stmt)
        user_configs = result.scalars().unique().all()

        validated_data = []
        for config in user_configs:
            model_name = "未知模型"
            if config.model and hasattr(config.model, "name"):
                model_name = config.model.name or "未知模型"
            
            platform_name = "未知平台"
            if config.model and config.model.platform and hasattr(config.model.platform, "name"):
                platform_name = config.model.platform.name or "未知平台"

            try:
                decrypted_key = decrypt_api_key(config.api_key) if config.api_key else ""
            except:
                decrypted_key = ""
            if len(decrypted_key) <= 8 and decrypted_key:
                api_key_mask = f"{decrypted_key[:4]}***"
            elif decrypted_key:
                api_key_mask = f"{decrypted_key[:4]}***{decrypted_key[-4:]}"
            else:
                api_key_mask = "***"

            validated_data.append({
                "id": config.id,
                "platform_name": platform_name,
                "model_name": model_name,
                "api_key_mask": api_key_mask,
                "base_url": config.base_url,
                "created_at": config.created_at
            })

        return [AIModelResponse(** data) for data in validated_data]

    except ValidationError as e:
        print(f"响应模型验证失败：{e}")
        raise HTTPException(status_code=500, detail="数据格式错误")
    except Exception as e:
        print(f"接口错误：{str(e)}")
        raise HTTPException(status_code=500, detail="获取配置失败")

# ----------------- 接口：获取支持的AI平台及模型列表 -----------------
@router.get("/platforms", response_model=List[PlatformModelResponse])
async def get_system_platforms(
    include_details: bool = Query(False, description="是否返回模型详情"),
    db: Session = Depends(get_db)
):
    platforms = db.query(Platform).filter(Platform.is_active == True).options(
        joinedload(Platform.models)
    ).order_by(Platform.name.asc()).all()
    
    return [platform.to_response(include_details=include_details) for platform in platforms]

@router.get("/platforms/{platform_id}/models", response_model=List[SystemModelResponse])
async def get_platform_models(
    platform_id: int,
    db: Session = Depends(get_db)
):
    platform = db.query(Platform).filter(
        Platform.id == platform_id,
        Platform.is_active == True
    ).first()
    if not platform:
        raise HTTPException(status_code=404, detail="系统未支持该平台或平台已禁用")
    
    models = db.query(Model).filter(
        Model.platform_id == platform_id,
        Model.is_supported == True
    ).order_by(Model.name.asc()).all()
    
    for model in models:
        model.platform_name = platform.name
    
    return models

# ----------------- 接口：新增AI模型配置 -----------------
@router.post("/keys", response_model=AIModelResponse)
async def create_ai_model(
    ai_model_create: AIModelCreate,  
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """用户新增AI模型配置"""
    # 1. 校验系统模型（内连接平台，确保平台存在且激活）
    system_model = db.query(Model).join(Platform).filter(
        Model.id == ai_model_create.model_id,
        Model.is_supported == True,
        Platform.is_active == True
    ).first()
    if not system_model:
        raise HTTPException(status_code=400, detail="系统未支持该模型或关联平台已禁用，无法创建配置")
    
    # 2. 校验重复配置
    existing_config = db.query(AIModel).filter(
        AIModel.user_id == current_user.id,
        AIModel.model_id == ai_model_create.model_id
    ).first()
    if existing_config:
        raise HTTPException(
            status_code=400, 
            detail=f"您已为「{system_model.platform.name}-{system_model.name}」配置过API Key，无需重复添加"
        )
    
    # 3. 加密API Key（添加异常捕获）
    try:
        encrypted_key = encrypt_api_key(ai_model_create.api_key)
    except Exception as e:
        raise HTTPException(
            status_code=400, 
            detail=f"API Key格式错误或加密失败：{str(e)}"
        )
    
    # 4. 创建新配置（显式设置时间戳）
    try:
        new_ai_config = AIModel(
            model_id=ai_model_create.model_id,
            user_id=current_user.id,
            api_key=encrypted_key,
            base_url=ai_model_create.base_url or system_model.platform.base_url,
            created_at=datetime.now(pytz.UTC),
            updated_at=datetime.now(pytz.UTC)
        )
        # 使用system_model赋值，避免访问未加载的关联
        new_ai_config.platform_name = system_model.platform.name
        new_ai_config.model_name = system_model.name
        
        db.add(new_ai_config)
        db.commit()
        db.refresh(new_ai_config)
        
        # 5. 处理API Key脱敏
        decrypted_key = decrypt_api_key(new_ai_config.api_key)
        if len(decrypted_key) <= 8:
            new_ai_config.api_key_mask = f"{decrypted_key[:4]}***"
        else:
            new_ai_config.api_key_mask = f"{decrypted_key[:4]}***{decrypted_key[-4:]}"
        
    except Exception as e:
        db.rollback()
        # 打印详细错误日志便于调试
        print(f"创建AI配置失败：{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"创建配置失败：{str(e)}")
    
    return new_ai_config

# ----------------- 接口：删除AI模型配置 -----------------
@router.delete("/keys/{key_id}")
async def delete_ai_model(
    key_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    ai_model = db.query(AIModel).filter(
        AIModel.id == key_id,
        AIModel.user_id == current_user.id
    ).options(
        joinedload(AIModel.model).joinedload(Model.platform)
    ).first()
    
    if not ai_model:
        raise HTTPException(status_code=404, detail="API配置不存在或无权删除")
    
    using_conversations = db.query(Conversation).filter(
        Conversation.user_id == current_user.id,
        Conversation.ai_model_id == key_id,
        Conversation.status == "active"
    ).first()
    
    if using_conversations:
        raise HTTPException(
            status_code=400, 
            detail=f"该模型正在被会话「{using_conversations.title}」使用，建议先修改会话的模型再删除"
        )

    try:
        db.delete(ai_model)
        db.commit()
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"删除AI模型配置失败：{str(e)}")

    platform_name = ai_model.model.platform.name
    model_name = ai_model.model.name
    return JSONResponse({"message": f"成功删除「{platform_name} - {model_name}」的API配置"})

# ----------------- 接口：编辑AI模型配置 -----------------
@router.put("/keys/{key_id}", response_model=AIModelResponse)
async def update_ai_model(
    key_id: int,
    ai_model_update: AIModelCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    ai_config = db.query(AIModel).filter(
        AIModel.id == key_id,
        AIModel.user_id == current_user.id
    ).options(
        joinedload(AIModel.model).joinedload(Model.platform)
    ).first()
    
    if not ai_config:
        raise HTTPException(status_code=404, detail="API配置不存在或无权编辑")

    system_model = db.query(Model).filter(
        Model.id == ai_model_update.model_id,
        Model.is_supported == True
    ).options(joinedload(Model.platform)).first()
    
    if not system_model:
        raise HTTPException(status_code=400, detail="系统未支持该模型，无法切换")

    if ai_model_update.model_id != ai_config.model_id:
        existing_config = db.query(AIModel).filter(
            AIModel.user_id == current_user.id,
            AIModel.model_id == ai_model_update.model_id
        ).first()
        if existing_config:
            raise HTTPException(
                status_code=400,
                detail=f"您已为「{system_model.platform.name}-{system_model.name}」配置过API Key，无需重复添加"
            )

    try:
        encrypted_key = encrypt_api_key(ai_model_update.api_key)
        ai_config.model_id = ai_model_update.model_id
        ai_config.api_key = encrypted_key
        ai_config.base_url = ai_model_update.base_url or system_model.platform.base_url
        ai_config.updated_at = datetime.now(pytz.UTC)

        db.commit()
        db.refresh(ai_config)

        try:
            decrypted_key = decrypt_api_key(ai_config.api_key)
            if len(decrypted_key) <= 8:
                ai_config.api_key_mask = f"{decrypted_key[:4]}***"
            else:
                ai_config.api_key_mask = f"{decrypted_key[:4]}***{decrypted_key[-4:]}"
        except Exception:
            ai_config.api_key_mask = "*******"

        ai_config.platform_name = ai_config.model.platform.name
        ai_config.model_name = ai_config.model.name

    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"更新AI配置失败：{str(e)}")

    return ai_config

# ----------------- 接口：删除模板 -----------------
@router.delete("/templates/{template_id}")
async def delete_template(
    template_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    template = db.query(Template).filter(
        Template.id == template_id,
        Template.user_id == current_user.id,
        Template.status == "active"
    ).first()
    
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在或无权删除")

    used_by_doc = db.query(DocumentHistory).filter(
        DocumentHistory.template_id == template_id,
        DocumentHistory.user_id == current_user.id
    ).first()
    
    if used_by_doc:
        raise HTTPException(
            status_code=400,
            detail=f"该模板已被公文「{used_by_doc.filename}」使用，无法删除"
        )

    try:
        template.status = "deleted"
        db.commit()
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"删除模板失败：{str(e)}")

    return JSONResponse({"message": f"成功删除模板「{template.original_name}」"})

# ----------------- 接口：测试文件写入权限 -----------------
@router.get("/test-write")
async def test_write():
    try:
        test_file = os.path.join(DOWNLOAD_DIR, "test_write.txt")
        with open(test_file, "w", encoding="utf-8") as f:
            f.write(f"测试写入权限 - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        return {"status": "success", "message": "generated_docs目录具备写入权限", "test_file": test_file}
    except Exception as e:
        return {"status": "error", "message": f"写入权限测试失败：{str(e)}"}

# ----------------- 补充接口：历史公文删除 -----------------
@router.delete("/history/{doc_id}")
async def delete_document_history(
    doc_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    doc = db.query(DocumentHistory).filter(
        DocumentHistory.id == doc_id,
        DocumentHistory.user_id == current_user.id
    ).first()
    
    if not doc:
        raise HTTPException(status_code=404, detail="历史公文不存在或无权删除")

    file_path = os.path.join(DOWNLOAD_DIR, doc.filename)
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
        except Exception as e:
            print(f"删除公文文件失败（{file_path}）：{str(e)}")

    try:
        db.delete(doc)
        db.commit()
        return JSONResponse({
            "message": f"历史公文「{doc.filename}」删除成功",
            "doc_id": doc_id
        })
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"删除公文记录失败：{str(e)}")

# ----------------- 补充接口：对话详情 -----------------
@router.get("/conversations/{conv_id}")
async def get_conversation_detail(
    conv_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    获取对话详情（含所有消息+关联的AI模型信息）
    - 权限：仅允许查看当前用户的对话
    - 用途：前端“继续对话”功能加载历史消息
    """
    # 1. 校验对话是否存在且属于当前用户（预加载关联的消息和AI模型）
    # 修正：将消息排序从joinedload移到单独的查询中
    conversation = db.query(Conversation).filter(
        Conversation.id == conv_id,
        Conversation.user_id == current_user.id
    ).options(
        joinedload(Conversation.messages),  # 仅加载消息，不在这里排序
        joinedload(Conversation.ai_model).joinedload(AIModel.model).joinedload(Model.platform)  # 预加载AI模型关联
    ).first()
    
    if not conversation:
        raise HTTPException(status_code=404, detail="对话不存在或无权访问")

    # 2. 单独对消息进行排序
    conversation.messages.sort(key=lambda x: x.created_at)

    # 3. 格式化响应（适配前端展示需求）
    return JSONResponse({
        "conv_id": conversation.id,
        "title": conversation.title,
        "created_at": conversation.created_at.astimezone(pytz.timezone("Asia/Shanghai")).strftime("%Y-%m-%d %H:%M:%S"),
        "updated_at": conversation.updated_at.astimezone(pytz.timezone("Asia/Shanghai")).strftime("%Y-%m-%d %H:%M:%S"),
        "used_model": {
            "platform": conversation.ai_model.model.platform.name,
            "model_name": conversation.ai_model.model.name,
            "config_id": conversation.ai_model.id
        },
        "messages": [
            {
                "id": msg.id,
                "role": msg.role,
                "content": msg.content,
                "docx_file": msg.docx_file,
                "created_at": msg.created_at.astimezone(pytz.timezone("Asia/Shanghai")).strftime("%Y-%m-%d %H:%M:%S")
            }
            for msg in conversation.messages
        ]
    })


# ----------------- 补充接口：对话删除 -----------------
@router.delete("/conversations/{conv_id}")
async def delete_conversation(
    conv_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    conversation = db.query(Conversation).filter(
        Conversation.id == conv_id,
        Conversation.user_id == current_user.id
    ).first()
    
    if not conversation:
        raise HTTPException(status_code=404, detail="对话不存在或无权删除")

    try:
        db.query(Message).filter(Message.conversation_id == conv_id).delete()
        conversation.status = "deleted"
        conversation.updated_at = datetime.now(pytz.UTC)
        db.commit()
        
        return JSONResponse({
            "message": f"对话「{conversation.title}」删除成功",
            "conv_id": conv_id
        })
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"删除对话失败：{str(e)}")

# ----------------- 补充接口：用户对话列表 -----------------
@router.get("/conversations")
async def get_user_conversations(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
    page: int = Query(1, ge=1),
    page_size: int = Query(10, ge=1, le=50),
    status: str = Query("active")
):
    query = db.query(Conversation).filter(
        Conversation.user_id == current_user.id,
        Conversation.status == status
    ).options(
        joinedload(Conversation.ai_model).joinedload(AIModel.model).joinedload(Model.platform)
    ).order_by(Conversation.updated_at.desc())

    total = query.count()
    conversations = query.offset((page - 1) * page_size).limit(page_size).all()

    result = []
    for conv in conversations:
        last_msg = db.query(Message).filter(
            Message.conversation_id == conv.id
        ).order_by(Message.created_at.desc()).first()
        
        result.append({
            "conv_id": conv.id,
            "title": conv.title,
            "created_at": conv.created_at.astimezone(pytz.timezone("Asia/Shanghai")).strftime("%Y-%m-%d %H:%M:%S"),
            "updated_at": conv.updated_at.astimezone(pytz.timezone("Asia/Shanghai")).strftime("%Y-%m-%d %H:%M:%S"),
            "used_model": f"{conv.ai_model.model.platform.name} - {conv.ai_model.model.name}",
            "last_message_preview": last_msg.content[:30] + "..." if last_msg else "无消息",
            "message_count": db.query(Message).filter(Message.conversation_id == conv.id).count()
        })

    return JSONResponse({
        "data": result,
        "pagination": {
            "total": total,
            "page": page,
            "page_size": page_size,
            "total_pages": (total + page_size - 1) // page_size
        }
    })

# ----------------- 公文历史详情 -----------------
@router.get("/history")
async def get_history(
    current_user = Depends(get_current_user), 
    db: Session = Depends(get_db),
    page: int = Query(1, ge=1),  # 页码，默认1
    page_size: int = Query(10, ge=1, le=50),  # 每页条数，限制1-50
    all: bool = Query(False, description="是否返回全量数据（忽略分页）")
):
    # 1. 基础查询：筛选当前用户的公文历史
    query = db.query(DocumentHistory).filter(DocumentHistory.user_id == current_user.id)
    total = query.count()
    
    # 2. 处理全量/分页逻辑
    if all:
        docs = query.order_by(DocumentHistory.created_at.desc()).all()
    else:
        docs = query.order_by(DocumentHistory.created_at.desc())\
            .offset((page - 1) * page_size)\
            .limit(page_size)\
            .all()
    
    # 3. 处理数据：补充模板名称预览，避免ellipsis
    result = []  # 初始化空数组（关键修复点）
    for doc in docs:
        # 查找公文关联的模板名称（若有）
        template_name = None
        if doc.template_id:
            template = db.query(Template).filter(
                Template.id == doc.template_id,
                Template.status == "active"  # 只关联未删除的模板
            ).first()
            template_name = template.original_name if template else "已删除模板"
        
        # 拼接单条公文历史数据
        result.append({
            "id": doc.id,
            "doc_type": doc.doc_type,
            "filename": doc.filename,
            "used_template": template_name,
            # 时间格式化：转为上海时区（适配前端展示）
            "created_at": doc.created_at.astimezone(pytz.timezone("Asia/Shanghai")).strftime("%Y-%m-%d %H:%M:%S"),
            # 内容预览：截取前50字符，超过则加省略号
            "content_preview": doc.content[:50] + "..." if len(doc.content) > 50 else doc.content
        })
    
    # 4. 构建响应（兼容全量/分页）
    response = {"data": result}
    if not all:
        response["pagination"] = {
            "total": total,
            "page": page,
            "page_size": page_size,
            "total_pages": (total + page_size - 1) // page_size  # 向上取整计算总页数
        }
    
    return JSONResponse(response)

# ----------------- 补充接口：用户通知 -----------------
@router.get("/notifications")
async def get_user_notifications(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
    unread_only: bool = Query(False)
):
    query = db.query(Notification).filter(Notification.user_id == current_user.id)
    if unread_only:
        query = query.filter(Notification.read == False)
    
    notifications = query.order_by(Notification.created_at.desc()).all()

    return JSONResponse({
        "total": len(notifications),
        "unread_count": db.query(Notification).filter(
            Notification.user_id == current_user.id,
            Notification.read == False
        ).count(),
        "notifications": [
            {
                "id": notify.id,
                "message": notify.message,
                "read": notify.read,
                "created_at": notify.created_at.astimezone(pytz.timezone("Asia/Shanghai")).strftime("%Y-%m-%d %H:%M:%S")
            }
            for notify in notifications
        ]
    })

# ----------------- 补充接口：标记通知为已读 -----------------
@router.patch("/notifications/{notify_id}/read")
async def mark_notification_read(
    notify_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    notify = db.query(Notification).filter(
        Notification.id == notify_id,
        Notification.user_id == current_user.id
    ).first()
    
    if not notify:
        raise HTTPException(status_code=404, detail="通知不存在或无权操作")

    try:
        notify.read = True
        db.commit()
        return JSONResponse({"message": "通知已标记为已读", "notify_id": notify_id})
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"标记通知失败：{str(e)}")
    