from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse, JSONResponse
import os
from .qwen_client import generate_text
from .utils import render_docx_from_template, save_uploaded_file
from dotenv import load_dotenv
import time
from docx import Document

load_dotenv()

router = APIRouter()
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), '..', 'uploads')
os.makedirs(UPLOAD_DIR, exist_ok=True)

PROMPTS = {
    "通知": "你是一个能写正式公文的助手，按照中国公文格式写一份{type}，注意语言正式、结构清晰，包含时间、地点、签发人等必要要素。",
    "请示": "你是一个能写正式公文的助手，写一份{type}，开门见山说明请求事由、请求事项、依据与建议。",
    "会议纪要": "写一份{type}，包含参会人员、时间地点、讨论要点、决议与后续事项。",
}

@router.post('/upload-template')
async def upload_template(file: UploadFile = File(...)):
    path = await save_uploaded_file(file, UPLOAD_DIR)
    return {"filename": os.path.basename(path), "path": path}

# 生成文件保存目录
DOWNLOAD_DIR = os.path.join(os.getcwd(), "generated_docs")
os.makedirs(DOWNLOAD_DIR, exist_ok=True)

@router.post('/generate')
async def generate_document(
    doc_type: str = Form(...),
    user_input: str = Form(...),
    template_filename: str | None = Form(None),
):
    base_prompt = PROMPTS.get(doc_type, "请写一份正式公文：{type}")
    prompt = base_prompt.format(type=doc_type) + "\n用户要求：" + user_input

    generated = generate_text(prompt)

    doc = Document()
    for line in generated.split("\n"):
        line = line.strip()
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        else:
            p = doc.add_paragraph()
            parts = line.split("**")
            for i, part in enumerate(parts):
                if i % 2 == 1:
                    run = p.add_run(part)
                    run.bold = True
                else:
                    subparts = part.split("*")
                    for j, sp in enumerate(subparts):
                        run = p.add_run(sp)
                        if j % 2 == 1:
                            run.italic = True

    filename = f"{doc_type}_{int(time.time())}.docx"
    file_path = os.path.join(DOWNLOAD_DIR, filename)
    doc.save(file_path)

    return JSONResponse({"text": generated, "filename": filename})


@router.get("/download/{filename}")
async def download(filename: str):
    file_path = os.path.join(DOWNLOAD_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件未找到")
    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )

