from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Depends
from fastapi.responses import FileResponse, JSONResponse
import os, time
from docx import Document
from dotenv import load_dotenv
from sqlalchemy.orm import Session
from .utils import load_data
from .qwen_client import generate_text
from .utils import save_uploaded_file
from .database import SessionLocal
from .models import DocumentHistory, Template
from .deps import get_current_user

load_dotenv()

router = APIRouter()

# 上传模板目录
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), '..', 'uploads')
os.makedirs(UPLOAD_DIR, exist_ok=True)
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CONVERSATIONS_FILE = os.path.join(PROJECT_ROOT, "conversations.json")

# 生成文件保存目录
DOWNLOAD_DIR = os.path.join(os.getcwd(), "generated_docs")
os.makedirs(DOWNLOAD_DIR, exist_ok=True)

# ----------------- PROMPTS -----------------
PROMPTS = {
    "通知": "你是一个能写正式公文的助手，按照中国公文格式写一份{type}，注意语言正式、结构清晰，包含时间、地点、签发人等必要要素。",
    "请示": "你是一个能写正式公文的助手，写一份{type}，开门见山说明请求事由、请求事项、依据与建议。",
    "会议纪要": "写一份{type}，包含参会人员、时间地点、讨论要点、决议与后续事项。",
}

# ----------------- 数据库 Session -----------------
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# ----------------- 上传模板 -----------------
@router.post("/upload-template")
async def upload_template(
    file: UploadFile = File(...),
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    path = await save_uploaded_file(file, UPLOAD_DIR)
    
    # 保存模板记录到数据库
    template = Template(
        user_id=current_user.id,
        filename=os.path.basename(path),
        original_name=file.filename
    )
    db.add(template)
    db.commit()
    db.refresh(template)
    
    return {"filename": template.filename, "original_name": template.original_name}

@router.post("/generate")
async def generate_document(
    doc_type: str = Form(...),
    user_input: str = Form(...),
    conv_id: int = Form(...),  
    template_content: str = Form(None),
    template_filename: str | None = Form(None),
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    base_prompt = PROMPTS.get(doc_type, "请写一份正式公文：{type}")
    prompt = base_prompt.format(type=doc_type) + "\n用户要求：" + user_input

    generated = generate_text(prompt)

    # 生成 DOCX
    doc = Document()
    for line in generated.split("\n"):
        line = line.strip()
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        else:
            p = doc.add_paragraph()
            parts = line.split("**")
            for i, part in enumerate(parts):
                if i % 2 == 1:
                    run = p.add_run(part)
                    run.bold = True
                else:
                    subparts = part.split("*")
                    for j, sp in enumerate(subparts):
                        run = p.add_run(sp)
                        if j % 2 == 1:
                            run.italic = True

    filename = f"{doc_type}_{int(time.time())}.docx"
    file_path = os.path.join(DOWNLOAD_DIR, filename)
    doc.save(file_path)

    # 保存到数据库
    doc_record = DocumentHistory(
        user_id=current_user.id,
        doc_type=doc_type,
        content=generated,
        filename=filename
    )
    db.add(doc_record)
    db.commit()

    # 🔁 把生成的公文加入对话
    data = load_data()
    for conv in data:
        if conv["id"] == conv_id:
            msg = {
                "id": int(time.time() * 1000),
                "role": "assistant",
                "content": generated,
                "docx_file": filename,
                "created_at": time.strftime("%Y-%m-%d %H:%M:%S")
            }
            conv["messages"].append(msg)
            save_data(data)
            break

    return JSONResponse({"text": generated, "filename": filename},
                        status_code=200)

# ----------------- 下载 DOCX -----------------
@router.get("/download/{filename}")
async def download(filename: str, current_user = Depends(get_current_user)):
    file_path = os.path.join(DOWNLOAD_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件未找到")
    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )

# ----------------- 获取用户文档历史 -----------------
@router.get("/history")
async def get_history(current_user = Depends(get_current_user), db: Session = Depends(get_db)):
    docs = db.query(DocumentHistory).filter(DocumentHistory.user_id == current_user.id).order_by(DocumentHistory.created_at.desc()).all()
    return [{"id": d.id, "doc_type": d.doc_type, "filename": d.filename, "created_at": d.created_at} for d in docs]

# ----------------- 获取用户模板列表 -----------------
@router.get("/templates")
async def get_templates(current_user = Depends(get_current_user), db: Session = Depends(get_db)):
    templates = db.query(Template).filter(Template.user_id == current_user.id).order_by(Template.uploaded_at.desc()).all()
    return [{"id": t.id, "filename": t.filename, "original_name": t.original_name, "uploaded_at": t.uploaded_at} for t in templates]


@router.get("/documents")
async def list_documents():
    files = os.listdir(DOWNLOAD_DIR)
    files = [f for f in files if f.endswith(".docx")]
    return {"documents": files}

@router.get("/template-content/{filename}")
async def get_template_content(filename: str, current_user=Depends(get_current_user)):
    file_path = os.path.join(UPLOAD_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="模板文件未找到")

    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return JSONResponse({"content": "\n".join(full_text)})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"读取模板失败: {str(e)}")
    
@router.route("/api/delete-conversation", methods=["POST"])
def delete_conversation():
    data = request.get_json()
    conv_id = data.get("id")

    if not conv_id:
        return jsonify({"error": "缺少会话ID"}), 400

    try:
        if os.path.exists(CONVERSATIONS_FILE):
            with open(CONVERSATIONS_FILE, "r", encoding="utf-8") as f:
                conversations = json.load(f)

            # 过滤掉指定ID的会话
            new_conversations = [c for c in conversations if c["id"] != conv_id]

            if len(conversations) == len(new_conversations):
                return jsonify({"error": "会话不存在"}), 404

            # 保存回文件
            with open(CONVERSATIONS_FILE, "w", encoding="utf-8") as f:
                json.dump(new_conversations, f, ensure_ascii=False, indent=2)

            return jsonify({"message": "删除成功"})

        return jsonify({"error": "会话文件不存在"}), 404

    except Exception as e:
        print(f"删除会话失败: {e}")
        return jsonify({"error": "删除失败"}), 500

def save_conversation(conv_id: int, user_msg: str, ai_msg: dict):
    # 加载现有会话
    if os.path.exists(CONVERSATIONS_FILE):
        with open(CONVERSATIONS_FILE, "r", encoding="utf-8") as f:
            conversations = json.load(f)
    else:
        conversations = []

    # 查找会话
    conv = next((c for c in conversations if c["id"] == conv_id), None)

    if not conv:
        conv = {
            "id": conv_id,
            "title": f"{doc_type} - {user_msg[:10]}...",
            "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            "updated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            "messages": []
        }
        conversations.append(conv)
    else:
        conv["updated_at"] = time.strftime("%Y-%m-%d %H:%M:%S")

    # 添加消息
    conv["messages"].append({
        "role": "user",
        "content": user_msg,
        "created_at": time.strftime("%Y-%m-%d %H:%M:%S")
    })
    conv["messages"].append({
        "role": "assistant",
        "content": ai_msg["text"],
        "docx_file": ai_msg["filename"],
        "created_at": time.strftime("%Y-%m-%d %H:%M:%S")
    })

    # 保存
    with open(CONVERSATIONS_FILE, "w", encoding="utf-8") as f:
        json.dump(conversations, f, ensure_ascii=False, indent=2)