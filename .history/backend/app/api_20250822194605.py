from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Depends
from fastapi.responses import FileResponse, JSONResponse
import os, time
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
from sqlalchemy.orm import Session
from .utils import load_data,save_data
from .qwen_client import generate_text
from .utils import save_uploaded_file
from .database import SessionLocal
from .models import DocumentHistory, Template
from .deps import get_current_user
import json
from fastapi import Request
from .conversations import _add_message_to_conversation, load_data, save_data
load_dotenv()

router = APIRouter()

# 上传模板目录
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), '..', 'uploads')
os.makedirs(UPLOAD_DIR, exist_ok=True)

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CONVERSATIONS_FILE = os.path.join(PROJECT_ROOT, "conversations.json")

# 生成文件保存目录
DOWNLOAD_DIR = os.path.join(os.getcwd(), "generated_docs")
os.makedirs(DOWNLOAD_DIR, exist_ok=True)

# ----------------- PROMPTS -----------------
PROMPTS = {
    "通知": "你是一个能写正式公文的助手，按照中国公文格式写一份{type}，注意语言正式、结构清晰，包含时间、地点、签发人等必要要素。",
    "请示": "你是一个能写正式公文的助手，写一份{type}，开门见山说明请求事由、请求事项、依据与建议。",
    "会议纪要": "写一份{type}，包含参会人员、时间地点、讨论要点、决议与后续事项。",
}

# ----------------- 数据库 Session -----------------
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# ----------------- 上传模板 -----------------
@router.post("/upload-template")
async def upload_template(
    file: UploadFile = File(...),
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    path = await save_uploaded_file(file, UPLOAD_DIR)
    
    # 保存模板记录到数据库
    template = Template(
        user_id=current_user.id,
        filename=os.path.basename(path),
        original_name=file.filename
    )
    db.add(template)
    db.commit()
    db.refresh(template)
    
    return {"filename": template.filename, "original_name": template.original_name}

python
浅色版本
from fastapi import APIRouter, Form, Depends, HTTPException, Response
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import time
from sqlalchemy.orm import Session
from fastapi.security import OAuth2PasswordBearer
from .auth import get_current_user
from .database import get_db
from .models import DocumentHistory

router = APIRouter()

DOWNLOAD_DIR = "downloads"

# 假设 PROMPTS 是一个字典，存储不同类型公文的提示信息
PROMPTS = {
    "通知": "请写一份正式的通知：",
    "请示": "请写一份正式的请示：",
    # 其他类型...
}

def generate_text(prompt):
    # 这里假设你有一个函数可以生成文本内容
    return "这是生成的公文内容\n\n# 标题\n\n这是正文内容。"

@router.post("/generate")
async def generate_document(
    doc_type: str = Form(...),
    user_input: str = Form(...),
    conv_id: int = Form(...),  
    template_content: str | None = Form(None),
    template_filename: str | None = Form(None),
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    base_prompt = PROMPTS.get(doc_type, "请写一份正式公文：{type}")
    prompt = base_prompt.format(type=doc_type) + "\n用户要求：" + user_input

    generated = generate_text(prompt)

    # 生成 DOCX
    doc = Document()
    
    # 设置全局字体（仿宋_GB2312）
    style = doc.styles['Normal']
    font = style.font
    font.name = 'FangSong_GB2312'  # 仿宋是公文标准
    font.size = Pt(16)  # 小四号字
    font.bold = False
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong_GB2312')  # 中文字体
    
    # 设置段落间距
    style.paragraph_format.space_before = Pt(0)  # 段前距为 0
    style.paragraph_format.space_after = Pt(0)   # 段后距为 0

    for line in generated.split("\n"):
        line = line.strip()
        if not line:
            continue  # 跳过空行
        
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        else:
            p = doc.add_paragraph()
            parts = line.split("**")
            for i, part in enumerate(parts):
                if i % 2 == 1:
                    run = p.add_run(part)
                    run.bold = True
                else:
                    subparts = part.split("*")
                    for j, sp in enumerate(subparts):
                        run = p.add_run(sp)
                        if j % 2 == 1:
                            run.italic = True

    filename = f"{doc_type}_{int(time.time())}.docx"
    file_path = os.path.join(DOWNLOAD_DIR, filename)
    doc.save(file_path)

    # 保存到数据库
    doc_record = DocumentHistory(
        user_id=current_user.id,
        doc_type=doc_type,
        content=generated,
        filename=filename
    )
    db.add(doc_record)
    db.commit()

    # 保存对话
    try:
        data = load_data()

        # 添加用户消息
        _add_message_to_conversation(data, conv_id, "user", user_input)

        # 添加 AI 消息（带 docx 文件名）
        _add_message_to_conversation(data, conv_id, "assistant", generated, docx_file=filename)

        # 保存到文件
        save_data(data)
    except Exception as e:
        print(f"保存对话失败: {e}")

    # 返回 JSON 响应
    return JSONResponse({
        "text": generated,
        "filename": filename,
        "html": markdown.markdown(generated, extensions=['fenced_code', 'tables'])  # 返回 HTML 预览
    }, status_code=200)

# ----------------- 下载 DOCX -----------------
@router.get("/download/{filename}")
async def download(filename: str, current_user = Depends(get_current_user)):
    file_path = os.path.join(DOWNLOAD_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件未找到")
    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )

# ----------------- 获取用户文档历史 -----------------
@router.get("/history")
async def get_history(current_user = Depends(get_current_user), db: Session = Depends(get_db)):
    docs = db.query(DocumentHistory).filter(DocumentHistory.user_id == current_user.id).order_by(DocumentHistory.created_at.desc()).all()
    return [{"id": d.id, "doc_type": d.doc_type, "filename": d.filename, "created_at": d.created_at} for d in docs]

# ----------------- 获取用户模板列表 -----------------
@router.get("/templates")
async def get_templates(current_user = Depends(get_current_user), db: Session = Depends(get_db)):
    templates = db.query(Template).filter(Template.user_id == current_user.id).order_by(Template.uploaded_at.desc()).all()
    return [{"id": t.id, "filename": t.filename, "original_name": t.original_name, "uploaded_at": t.uploaded_at} for t in templates]


@router.get("/documents")
async def list_documents():
    files = os.listdir(DOWNLOAD_DIR)
    files = [f for f in files if f.endswith(".docx")]
    return {"documents": files}

@router.get("/template-content/{filename}")
async def get_template_content(filename: str, current_user=Depends(get_current_user)):
    file_path = os.path.join(UPLOAD_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="模板文件未找到")

    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return JSONResponse({"content": "\n".join(full_text)})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"读取模板失败: {str(e)}")
    

@router.post("/delete-conversation")
async def delete_conversation(request: Request):
    data = await request.json()
    conv_id = data.get("id")

    if not conv_id:
        raise HTTPException(status_code=400, detail="缺少会话ID")

    try:
        if os.path.exists(CONVERSATIONS_FILE):
            with open(CONVERSATIONS_FILE, "r", encoding="utf-8") as f:
                conversations = json.load(f)

            new_conversations = [c for c in conversations if c["id"] != conv_id]

            if len(conversations) == len(new_conversations):
                raise HTTPException(status_code=404, detail="会话不存在")

            with open(CONVERSATIONS_FILE, "w", encoding="utf-8") as f:
                json.dump(new_conversations, f, ensure_ascii=False, indent=2)

            return JSONResponse({"message": "删除成功"})
        else:
            raise HTTPException(status_code=404, detail="会话文件不存在")

    except Exception as e:
        print(f"删除会话失败: {e}")
        raise HTTPException(status_code=500, detail="删除失败")
@router.get("/test-write")
async def test_write():
    try:
        test_file = os.path.join(DOWNLOAD_DIR, "test_write.txt")
        with open(test_file, "w", encoding="utf-8") as f:
            f.write("测试写入权限\n")
        return {"status": "success", "message": "可以写入 generated_docs"}
    except Exception as e:
        return {"status": "error", "message": str(e)}

def save_conversation(conv_id: int, user_msg: str, ai_msg: dict):
    # 加载现有会话
    if os.path.exists(CONVERSATIONS_FILE):
        with open(CONVERSATIONS_FILE, "r", encoding="utf-8") as f:
            conversations = json.load(f)
    else:
        conversations = []

    # 查找会话
    conv = next((c for c in conversations if c["id"] == conv_id), None)

    # 如果是新会话，创建一个
    if not conv:
        # 从 user_msg 提取标题（前10个字）
        title = user_msg.strip()[:10] + "..." if len(user_msg.strip()) > 10 else user_msg.strip()
        conv = {
            "id": conv_id,
            "title": f"{title}",
            "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            "updated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            "messages": []
        }
        conversations.append(conv)
    else:
        conv["updated_at"] = time.strftime("%Y-%m-%d %H:%M:%S")

    # 添加用户消息
    conv["messages"].append({
        "role": "user",
        "content": user_msg,
        "created_at": time.strftime("%Y-%m-%d %H:%M:%S")
    })

    # 添加 AI 消息
    conv["messages"].append({
        "role": "assistant",
        "content": ai_msg["text"],
        "docx_file": ai_msg["filename"],
        "created_at": time.strftime("%Y-%m-%d %H:%M:%S")
    })

    # 保存回文件
    with open(CONVERSATIONS_FILE, "w", encoding="utf-8") as f:
        json.dump(conversations, f, ensure_ascii=False, indent=2)