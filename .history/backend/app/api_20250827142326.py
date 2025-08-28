# api.py
from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Depends, Query
from fastapi.responses import FileResponse, JSONResponse
import os, time
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
from sqlalchemy.orm import Session, joinedload
from sqlalchemy.future import select
from sqlalchemy.ext.asyncio import AsyncSession
from pydantic import ValidationError
from typing import List, Optional
from datetime import datetime
import pytz
import traceback
import urllib
# 导入自定义模块（确保路径正确）
from .utils import save_uploaded_file
from .AI_client import generate_text_for_user, get_user_model_preference  # 多平台Client核心函数
from .database import SessionLocal, get_db,get_async_db 
from .models import (
    DocumentHistory, Template, AIModel, Conversation, Message, Platform,
    AIModelResponse, PlatformModelResponse, User, AIModelCreate, SystemModelResponse, Model
)
from .deps import get_current_user
from .encryption import encrypt_api_key, decrypt_api_key

# 加载环境变量
load_dotenv()

router = APIRouter()

# ----------------- 目录配置（保持原路径逻辑） -----------------
# 上传模板目录
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), '..', 'uploads')
os.makedirs(UPLOAD_DIR, exist_ok=True)

# 生成文件保存目录
DOWNLOAD_DIR = os.path.join(os.getcwd(), "generated_docs")
os.makedirs(DOWNLOAD_DIR, exist_ok=True)

# ----------------- 公文生成Prompt配置（保持原逻辑） -----------------
PROMPTS = {
    "通知": "你是专业的正式公文写作助手，需严格按照《党政机关公文处理工作条例》规定的通知格式撰写。内容需包含标题（发文机关+事由+文种）、主送机关、正文（明确通知缘由、具体事项、执行要求，涉及任务需标注责任主体）、附件说明（如有附件需注明附件名称及数量）、发文机关署名、成文日期、签发人（上行文需标注），语言需庄重严谨、逻辑清晰，避免歧义。",
    "请示": "你是专业的正式公文写作助手，需严格遵循请示“一文一事、单头主送”的核心要求撰写。内容需包含标题（发文机关+事由+文种）、主送机关（仅限一个直接上级机关）、正文（开门见山说明请示缘由，需引用政策依据或实际需求；清晰列出请示事项，做到具体明确；提出合理可行的建议方案供上级决策）、发文机关署名、成文日期、签发人，语言需谦恭得体，避免多头请示或夹带报告事项。",
    "会议纪要": "你是专业的正式公文写作助手，需按规范会议纪要格式撰写，确保内容客观准确、要素完整。需包含标题（会议名称+文种，如“XX会议纪要”）、基本信息（会议时间、会议地点、主持人、记录人、参会人员/列席人员/缺席人员名单及职务）、正文（分点梳理会议讨论要点，需准确反映不同观点；明确会议形成的决议事项，标注每项决议的责任单位、完成时限）、结尾（如需补充说明未尽事宜或后续沟通方式，可简要标注），语言需简洁凝练，避免主观评价。",
    "报告": "你是专业的正式公文写作助手，需严格按照报告的法定格式与写作要求撰写，区分“汇报工作”“反映情况”“答复询问”等不同报告类型。内容需包含标题（发文机关+事由+文种，如“XX关于XX工作的报告”）、主送机关、正文（开头简要说明报告背景或依据，主体部分详细阐述工作进展、成果、存在问题、原因分析及下一步计划，答复类报告需直接回应上级询问事项）、发文机关署名、成文日期、签发人（上行文需标注），语言需客观真实，不得夹带请示事项。",
    "函": "你是专业的正式公文写作助手，需根据函的“不相隶属机关商洽工作、询问答复问题、请求批准事项”功能，按规范格式撰写。内容需包含标题（发文机关+事由+文种，如“XX关于XX事宜的函”或“XX关于XX问题的复函”）、主送机关（明确的不相隶属机关）、正文（商洽类需说明事由及具体需求，询问类需列出清晰的问题清单，答复类需直接回应来函事项，请求批准类需说明依据及具体请求）、发文机关署名、成文日期，语言需平和得体，避免命令性或指令性表述。",
    "对话": "你是专业的公文场景对话撰写助手，需围绕公文办理全流程（如公文起草沟通、审批意见反馈、事项协调对接等）撰写对话内容。需明确对话场景（如“公文起草小组沟通会议对话”“上下级机关审批意见反馈对话”）、对话主体（标注角色及职务，如“起草人-XX部科员”“审批人-XX局副局长”）、对话逻辑（需符合公文办理规范，内容需聚焦具体事项，如格式修改建议、内容补充要求、办理时限确认等），语言需贴合职场沟通语境，既保持正式性，又体现沟通的针对性与高效性。"
}

# ----------------- 核心接口：上传模板（补充模板内容存储） -----------------
@router.post("/upload-template")
async def upload_template(
    file: UploadFile = File(...),
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # 保存模板文件到本地目录
    file_path = await save_uploaded_file(file, UPLOAD_DIR)
    filename = os.path.basename(file_path)
    
    # 读取模板内容（存入数据库，避免后续重复读取文件）
    template_content = ""
    try:
        doc = Document(file_path)
        template_content = "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"读取模板内容失败：{str(e)}")
    
    # 保存模板记录到数据库（含内容）
    template = Template(
        user_id=current_user.id,
        filename=filename,
        original_name=file.filename,
        content=template_content,  # 存储模板内容，提升后续生成效率
        status="active"  # 补充models中Template必需的status字段
    )
    db.add(template)
    db.commit()
    db.refresh(template)
    
    return {
        "id": template.id,
        "filename": template.filename,
        "original_name": template.original_name,
        "uploaded_at": template.uploaded_at
    }

# ----------------- 核心接口：公文生成（适配多平台Client+数据库会话） -----------------
@router.post("/generate")
async def generate_document(
    doc_type: str = Form(...),
    user_input: str = Form(...),
    conv_id: Optional[int] = Form(None),  # 可选：已有会话ID
    ai_model_id: Optional[int] = Form(None),  # 可选：用户手动选择的模型ID
    template_id: Optional[int] = Form(None),  # 改用模板ID（替代原filename，避免文件依赖）
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # 1. 组装生成Prompt（含公文类型专属System Prompt）
    base_prompt = PROMPTS.get(doc_type, f"请写一份正式公文：{doc_type}")
    # 若有模板，从数据库获取模板内容补充到Prompt
    if template_id:
        template = db.query(Template).filter(
            Template.id == template_id,
            Template.user_id == current_user.id,
            Template.status == "active"  # 只使用有效模板
        ).first()
        if not template:
            raise HTTPException(status_code=404, detail="指定模板不存在或无权访问")
        prompt = f"{base_prompt}\n模板内容：{template.content}\n用户要求：{user_input}"
    else:
        prompt = f"{base_prompt}\n用户要求：{user_input}"

    # 2. 获取用户AI模型配置（优先手动选择，其次默认偏好）
    user_config = None
    selected_ai_model = None
    
    if ai_model_id:
        # 优先使用用户手动选择的模型（校验归属+预加载关联）
        selected_ai_model = db.query(AIModel).filter(
            AIModel.id == ai_model_id,
            AIModel.user_id == current_user.id
        ).options(joinedload(AIModel.model)).first()  # 预加载Model关联
        if not selected_ai_model:
            raise HTTPException(status_code=404, detail="选择的AI模型不存在或无权访问")
        
        # ---------------- 关键修改：平台名称映射 ----------------
        # 1. 获取数据库中的平台名称（如 "Alibaba"）
        db_platform_name = selected_ai_model.model.platform.name
        # 2. 建立映射关系：数据库名称 → 客户端工厂支持的标识
        platform_mapping = {
            "Alibaba": "qwen",    # 核心映射：Alibaba → qwen（通义千问）
            "OpenAI": "openai",   # 其他平台按需补充（确保与数据库一致）
            "Anthropic": "anthropic",
            "Google": "gemini"
        }
        # 3. 转换为工厂类支持的平台标识（默认用原名称，避免未知平台报错）
        provider = platform_mapping.get(db_platform_name, db_platform_name.lower())
        
        # 组装用户配置（使用转换后的provider）
        user_config = {
            "provider": provider,  # 改用映射后的标识（如 "qwen"）
            "api_key": decrypt_api_key(selected_ai_model.api_key),  # 解密API Key
            "model_name": selected_ai_model.model.name,  # 从Model获取模型名
            "base_url": selected_ai_model.effective_base_url  # 使用models中定义的property
        }
    else:
        # 无手动选择时，用用户默认偏好（上次使用/最新添加）
        user_config = get_user_model_preference(user_id=current_user.id, db=db)
        if not user_config:
            raise HTTPException(status_code=400, detail="请先在「API Keys管理」中添加AI模型配置")
        
        # ---------------- 关键修改：默认偏好也需映射 ----------------
        # 对默认偏好的平台名称进行同样映射
        platform_mapping = {
            "Alibaba": "qwen",
            "OpenAI": "openai",
            "Anthropic": "anthropic",
            "Google": "gemini"
        }
        user_config["provider"] = platform_mapping.get(
            user_config["provider"], 
            user_config["provider"].lower()
        )
        
        # 同步获取模型实例（用于后续关联会话）
        selected_ai_model = db.query(AIModel).filter(
            AIModel.model.has(
                Model.platform.has(name=user_config["provider"]),  # 注意：此处需与映射后的名称匹配
                name=user_config["model_name"]
            ),
            AIModel.user_id == current_user.id
        ).options(joinedload(AIModel.model)).first()
        
        if not selected_ai_model:
            raise HTTPException(status_code=500, detail="用户默认模型配置异常，请重新添加")
    # 3. 调用多平台Client生成公文内容
    try:
        generated = generate_text_for_user(
            user_id=current_user.id,
            prompt=prompt,
            db=db,
            system_prompt=base_prompt  # 传入公文类型专属System Prompt
        )
    except Exception as e:
        # 捕获API Key错误并转换提示
        error_msg = str(e)
        if "401" in error_msg and "invalid_api_key" in error_msg.lower():
            raise HTTPException(
                status_code=401, 
                detail="API Key不正确或已失效，请检查并重新配置"
            )
        elif "不支持的平台" in error_msg:
            raise HTTPException(
                status_code=400,
                detail="所选AI平台暂不支持，请更换其他平台"
            )
        else:
            raise HTTPException(
                status_code=500, 
                detail=f"AI生成公文失败：{error_msg}"
            )

    # 4. 生成DOCX文件（保持原格式逻辑）
    doc = Document()
    # 设置公文标准格式（仿宋_GB2312、小四号字）
    style = doc.styles['Normal']
    font = style.font
    font.name = 'FangSong_GB2312'
    font.size = Pt(16)
    font.bold = False
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong_GB2312')
    # 设置段落间距（公文标准：段前/段后0，行距固定28磅，此处简化为段间距0）
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # 解析生成内容，构建DOCX结构
    for line in generated.split("\n"):
        line = line.strip()
        if not line:
            continue
        # 处理标题（# 一级标题、## 二级标题等）
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        else:
            # 处理加粗/斜体（Markdown格式）
            p = doc.add_paragraph()
            parts = line.split("**")
            for i, part in enumerate(parts):
                if i % 2 == 1:
                    run = p.add_run(part)
                    run.bold = True
                else:
                    subparts = part.split("*")
                    for j, sp in enumerate(subparts):
                        run = p.add_run(sp)
                        if j % 2 == 1:
                            run.italic = True

    # 保存DOCX文件（加用户ID避免文件名冲突）
    filename = f"{doc_type}_{current_user.id}_{int(time.time())}.docx"
    file_path = os.path.join(DOWNLOAD_DIR, filename)
    doc.save(file_path)

    # 5. 保存公文记录到DocumentHistory
    doc_record = DocumentHistory(
        user_id=current_user.id,
        doc_type=doc_type,
        content=generated,
        filename=filename,
        template_id=template_id  # 关联模板ID（而非文件名）
    )
    db.add(doc_record)
    db.commit()
    db.refresh(doc_record)

    # 6. 关联数据库会话（创建/更新Conversation，移除models中不存在的content字段）
    if conv_id:
        # 已有会话：更新会话时间和关联模型
        conversation = db.query(Conversation).filter(
            Conversation.id == conv_id,
            Conversation.user_id == current_user.id
        ).first()
        if not conversation:
            raise HTTPException(status_code=404, detail="指定会话不存在或无权访问")
        conversation.updated_at = datetime.now(pytz.UTC)
        conversation.ai_model_id = selected_ai_model.id  # 更新会话关联的模型
        conversation.status = "active"
    else:
        # 新会话：创建会话（自动生成标题，移除content字段）
        title = f"{doc_type}生成_{user_input.strip()[:8]}..." if len(user_input.strip()) > 8 else f"{doc_type}生成"
        conversation = Conversation(
            user_id=current_user.id,
            ai_model_id=selected_ai_model.id,  # 关联当前使用的模型
            title=title,
            status="active",  # 补充models中必需的status字段
            created_at=datetime.now(pytz.UTC),
            updated_at=datetime.now(pytz.UTC)
        )
        db.add(conversation)
    db.commit()
    db.refresh(conversation)

    # 7. 保存消息到Message表（用户输入+AI回复）
    user_message = Message(
        conversation_id=conversation.id,
        role="user",
        content=user_input,
        created_at=datetime.now(pytz.UTC)
    )
    ai_message = Message(
        conversation_id=conversation.id,
        role="assistant",
        content=generated,
        docx_file=filename,  # 关联生成的DOCX文件
        created_at=datetime.now(pytz.UTC)
    )
    db.add_all([user_message, ai_message])
    db.commit()

    # 8. 返回结果（含会话ID，便于前端后续继续对话）
    return JSONResponse({
        "text": generated,
        "filename": filename,
        "conv_id": conversation.id,  # 返回会话ID，支持继续对话
        "doc_id": doc_record.id,
        # 显示完整的平台+模型名称（从关联表获取）
        "used_model": f"{selected_ai_model.model.platform.name} - {selected_ai_model.model.name}"
    }, status_code=200)

# ----------------- 接口：下载生成的DOCX文件（补充DB依赖+权限校验） -----------------
@router.get("/download/{filename}")
async def download(
    filename: str, 
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # 1. 解码URL编码的文件名（关键：前端传递的是编码后的文件名）
    decoded_filename = urllib.parse.unquote(filename)
    
    # 2. 校验文件是否存在（使用解码后的文件名）
    file_path = os.path.join(DOWNLOAD_DIR, decoded_filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件未找到或已被删除")
    
    # 3. 校验文件归属（使用解码后的文件名查询数据库）
    doc_record = db.query(DocumentHistory).filter(
        DocumentHistory.filename == decoded_filename,  # 与数据库原始文件名匹配
        DocumentHistory.user_id == current_user.id
    ).first()
    if not doc_record:
        raise HTTPException(status_code=403, detail="无权访问此文件（非本人生成）")
    
    # 4. 对文件名进行HTTP响应头编码（支持中文等特殊字符）
    encoded_filename = urllib.parse.quote(decoded_filename)
    
    # 5. 返回文件下载（使用标准的UTF-8编码格式）
    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            # 标准写法：明确指定UTF-8编码
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
        }
    )
    

# ----------------- 接口：获取用户公文历史 -----------------
@router.get("/history")
async def get_history(
    current_user = Depends(get_current_user), 
    db: Session = Depends(get_db)
):
    # 查询当前用户的公文历史（按生成时间倒序）
    docs = db.query(DocumentHistory).filter(
        DocumentHistory.user_id == current_user.id
    ).order_by(DocumentHistory.created_at.desc()).all()
    
    # 补充关联的模板名称（若有）
    result = []
    for doc in docs:
        template_name = None
        if doc.template_id:
            template = db.query(Template).filter(
                Template.id == doc.template_id,
                Template.status == "active"
            ).first()
            template_name = template.original_name if template else None
        
        result.append({
            "id": doc.id,
            "doc_type": doc.doc_type,
            "filename": doc.filename,
            "used_template": template_name,
            "created_at": doc.created_at.strftime("%Y-%m-%d %H:%M:%S"),  # 格式化时间，便于前端展示
            "content_preview": doc.content[:50] + "..." if len(doc.content) > 50 else doc.content  # 内容预览
        })
    return result

# ----------------- 接口：获取用户模板列表 -----------------
@router.get("/templates")
async def get_templates(
    current_user = Depends(get_current_user), 
    db: Session = Depends(get_db)
):
    # 查询当前用户的有效模板（按上传时间倒序）
    templates = db.query(Template).filter(
        Template.user_id == current_user.id,
        Template.status == "active"  # 只返回有效模板
    ).order_by(Template.uploaded_at.desc()).all()
    
    return [
        {
            "id": t.id,
            "filename": t.filename,
            "original_name": t.original_name,
            "uploaded_at": t.uploaded_at.strftime("%Y-%m-%d %H:%M:%S"),
            "content_preview": t.content[:30] + "..." if len(t.content) > 30 else t.content  # 模板内容预览
        }
        for t in templates
    ]

# ----------------- 接口：获取模板内容（从数据库读取，替代文件读取） -----------------
@router.get("/template-content/{template_id}")  # 改用template_id，避免文件名依赖
async def get_template_content(
    template_id: int, 
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # 从数据库查询有效模板（含内容），避免重复读取文件
    template = db.query(Template).filter(
        Template.id == template_id,
        Template.user_id == current_user.id,
        Template.status == "active"
    ).first()
    
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在或无权访问")
    
    # 直接返回数据库中的内容（无需再读取文件）
    return JSONResponse({
        "id": template.id,
        "original_name": template.original_name,
        "content": template.content
    })

# ----------------- 接口：获取用户API Key列表（解密脱敏） -----------------
@router.get("/keys", response_model=List[AIModelResponse])
async def get_user_ai_configs(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_async_db)
):
    try:
        # 1. 查询并预加载关联
        stmt = select(AIModel).filter(
            AIModel.user_id == current_user.id
        ).options(
            joinedload(AIModel.model).joinedload(Model.platform)
        ).order_by(AIModel.created_at.desc())
        
        result = await db.execute(stmt)
        user_configs = result.scalars().unique().all()

        # 2. 手动提取字段值并构造验证数据（关键步骤）
        validated_data = []
        for config in user_configs:
            # 手动提取 model.name（确保非空）
            model_name = "未知模型"
            if config.model and hasattr(config.model, "name"):
                model_name = config.model.name or "未知模型"  # 处理空字符串
            
            # 手动提取 model.platform.name（确保非空）
            platform_name = "未知平台"
            if config.model and config.model.platform and hasattr(config.model.platform, "name"):
                platform_name = config.model.platform.name or "未知平台"  # 处理空字符串

            # 脱敏 API Key
            try:
                decrypted_key = decrypt_api_key(config.api_key) if config.api_key else ""
            except:
                decrypted_key = ""
            if len(decrypted_key) <= 8 and decrypted_key:
                api_key_mask = f"{decrypted_key[:4]}***"
            elif decrypted_key:
                api_key_mask = f"{decrypted_key[:4]}***{decrypted_key[-4:]}"
            else:
                api_key_mask = "***"

            # 构造符合 AIModelResponse 结构的字典
            validated_data.append({
                "id": config.id,
                "platform_name": platform_name,  # 使用手动提取的值
                "model_name": model_name,        # 使用手动提取的值
                "api_key_mask": api_key_mask,
                "base_url": config.base_url,
                "created_at": config.created_at
            })

        # 3. 用 Pydantic 验证构造的数据（此时字段均已存在）
        return [AIModelResponse(**data) for data in validated_data]

    except ValidationError as e:
        print(f"响应模型验证失败：{e}")
        raise HTTPException(status_code=500, detail="数据格式错误")
    except Exception as e:
        print(f"接口错误：{str(e)}")
        raise HTTPException(status_code=500, detail="获取配置失败")
    
# ----------------- 接口：获取支持的AI平台及模型列表 -----------------
@router.get("/platforms", response_model=List[PlatformModelResponse])
async def get_system_platforms(
    include_details: bool = Query(False, description="是否返回模型详情（如描述、是否支持）"),
    db: Session = Depends(get_db)
):
    """
    获取系统支持的AI平台列表（从数据库 Platform/Model 表查询，替代原硬编码）
    - include_details=True：返回模型详情（供前端选择模型时展示描述）
    - include_details=False：仅返回平台名称和模型名称列表（轻量查询）
    """
    # 只查询系统启用的平台（is_active=True），预加载关联的模型
    platforms = db.query(Platform).filter(Platform.is_active == True).options(
        joinedload(Platform.models)  # 预加载平台下的模型
    ).order_by(Platform.name.asc()).all()
    
    # 调用 Platform.to_response() 转换为 Pydantic 响应格式
    return [platform.to_response(include_details=include_details) for platform in platforms]


@router.get("/platforms/{platform_id}/models", response_model=List[SystemModelResponse])
async def get_platform_models(
    platform_id: int,
    db: Session = Depends(get_db)
):
    """
    获取指定平台下的所有系统模型（供前端“按平台筛选模型”）
    """
    # 1. 校验平台是否存在且启用
    platform = db.query(Platform).filter(
        Platform.id == platform_id,
        Platform.is_active == True
    ).first()
    if not platform:
        raise HTTPException(status_code=404, detail="系统未支持该平台或平台已禁用")
    
    # 2. 查询该平台下的所有支持模型（is_supported=True）
    models = db.query(Model).filter(
        Model.platform_id == platform_id,
        Model.is_supported == True
    ).order_by(Model.name.asc()).all()
    
    # 补充平台名称（SystemModelResponse需要platform_name字段）
    for model in models:
        model.platform_name = platform.name
    
    return models

# ----------------- 接口：新增AI模型配置（加密存储） -----------------
@router.post("/keys", response_model=AIModelResponse)
async def create_ai_model(
    ai_model_create: AIModelCreate,  
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    用户新增AI模型配置（基于系统模型ID，不再手动输入平台/模型名称）
    """
    # 1. 校验系统模型是否存在且支持
    system_model = db.query(Model).filter(
        Model.id == ai_model_create.model_id,
        Model.is_supported == True
    ).options(joinedload(Model.platform)).first()  # 预加载Platform关联
    if not system_model:
        raise HTTPException(status_code=400, detail="系统未支持该模型，无法创建配置")
    
    # 2. 校验用户是否已为该系统模型配置过API Key（避免重复）
    existing_config = db.query(AIModel).filter(
        AIModel.user_id == current_user.id,
        AIModel.model_id == ai_model_create.model_id
    ).first()
    if existing_config:
        raise HTTPException(
            status_code=400, 
            detail=f"您已为「{system_model.platform.name}-{system_model.name}」配置过API Key，无需重复添加"
        )
    
    # 3. 加密API Key并存储（复用原有加密逻辑）
    try:
        encrypted_key = encrypt_api_key(ai_model_create.api_key)
        new_ai_config = AIModel(
            model_id=ai_model_create.model_id,
            user_id=current_user.id,
            api_key=encrypted_key,
            # 复用系统平台默认BaseURL（用户未传时）
            base_url=ai_model_create.base_url or system_model.platform.base_url
        )
        new_ai_config.platform_name = new_ai_config.model.platform.name
        new_ai_config.model_name = new_ai_config.model.name
        db.add(new_ai_config)
        db.commit()
        db.refresh(new_ai_config)
        
        # 手动处理API Key脱敏（供响应返回）
        decrypted_key = decrypt_api_key(new_ai_config.api_key)
        if len(decrypted_key) <= 8:
            new_ai_config.api_key_mask = f"{decrypted_key[:4]}***"
        else:
            new_ai_config.api_key_mask = f"{decrypted_key[:4]}***{decrypted_key[-4:]}"
        
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"创建AI配置失败：{str(e)}")
    
    return new_ai_config

# ----------------- 接口：删除AI模型配置 -----------------
@router.delete("/keys/{key_id}")
async def delete_ai_model(
    key_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    # 1. 查询待删除的模型（校验归属+预加载关联）
    ai_model = db.query(AIModel).filter(
        AIModel.id == key_id,
        AIModel.user_id == current_user.id
    ).options(
        joinedload(AIModel.model).joinedload(Model.platform)
    ).first()
    
    if not ai_model:
        raise HTTPException(status_code=404, detail="API配置不存在或无权删除")
    
    # 2. 校验是否有会话正在使用该模型（避免删除正在使用的模型）
    using_conversations = db.query(Conversation).filter(
        Conversation.user_id == current_user.id,
        Conversation.ai_model_id == key_id,
        Conversation.status == "active"
    ).first()
    
    if using_conversations:
        raise HTTPException(
            status_code=400, 
            detail=f"该模型正在被会话「{using_conversations.title}」使用，建议先修改会话的模型再删除"
        )

    # 3. 执行删除
    try:
        db.delete(ai_model)
        db.commit()
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"删除AI模型配置失败：{str(e)}")

    # 返回完整的平台+模型名称（从关联表获取）
    platform_name = ai_model.model.platform.name
    model_name = ai_model.model.name
    return JSONResponse({"message": f"成功删除「{platform_name} - {model_name}」的API配置"})
# ----------------- 接口：编辑AI模型配置（新增PUT接口） -----------------
@router.put("/keys/{key_id}", response_model=AIModelResponse)
async def update_ai_model(
    key_id: int,  # 待编辑的配置ID
    ai_model_update: AIModelCreate,  # 复用创建配置的请求模型（model_id/api_key/base_url）
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    用户编辑已有的AI模型配置（支持修改API Key、BaseURL、切换同平台/跨平台模型）
    """
    # 1. 校验待编辑的配置是否存在，且属于当前用户（预加载关联表，避免后续N+1查询）
    ai_config = db.query(AIModel).filter(
        AIModel.id == key_id,
        AIModel.user_id == current_user.id
    ).options(
        joinedload(AIModel.model).joinedload(Model.platform)  # 预加载Model→Platform关联
    ).first()
    
    if not ai_config:
        raise HTTPException(status_code=404, detail="API配置不存在或无权编辑")

    # 2. 校验用户要切换的新模型（model_id）是否为系统支持的模型
    system_model = db.query(Model).filter(
        Model.id == ai_model_update.model_id,
        Model.is_supported == True  # 仅允许切换到系统支持的模型
    ).options(joinedload(Model.platform)).first()
    
    if not system_model:
        raise HTTPException(status_code=400, detail="系统未支持该模型，无法切换")

    # 3. 校验是否重复：若用户切换了模型，需检查是否已为新模型配置过API Key
    if ai_model_update.model_id != ai_config.model_id:  # 仅当模型ID变化时才校验重复
        existing_config = db.query(AIModel).filter(
            AIModel.user_id == current_user.id,
            AIModel.model_id == ai_model_update.model_id
        ).first()
        if existing_config:
            raise HTTPException(
                status_code=400,
                detail=f"您已为「{system_model.platform.name}-{system_model.name}」配置过API Key，无需重复添加"
            )

    # 4. 执行更新逻辑（加密新API Key、处理BaseURL默认值、更新时间戳）
    try:
        # 加密新提交的明文API Key
        encrypted_key = encrypt_api_key(ai_model_update.api_key)
        # 更新核心字段
        ai_config.model_id = ai_model_update.model_id  # 切换模型（若有变化）
        ai_config.api_key = encrypted_key  # 更新加密后的API Key
        # BaseURL：用户传了就用用户的，没传就用新模型所属平台的默认BaseURL
        ai_config.base_url = ai_model_update.base_url or system_model.platform.base_url
        ai_config.updated_at = datetime.now(pytz.UTC)  # 更新时间戳

        # 5. 提交事务（核心逻辑：仅更新数据库，脱敏逻辑后续处理，避免事务回滚）
        db.commit()
        db.refresh(ai_config)  # 刷新实例，获取最新关联数据

        # 6. 处理API Key脱敏（供响应返回，失败不影响数据库更新）
        try:
            decrypted_key = decrypt_api_key(ai_config.api_key)
            if len(decrypted_key) <= 8:
                ai_config.api_key_mask = f"{decrypted_key[:4]}***"
            else:
                ai_config.api_key_mask = f"{decrypted_key[:4]}***{decrypted_key[-4:]}"
        except Exception:
            ai_config.api_key_mask = "*******"  # 脱敏失败时显示默认值

        # 7. 补充响应所需的关联字段（platform_name/model_name，适配AIModelResponse）
        ai_config.platform_name = ai_config.model.platform.name
        ai_config.model_name = ai_config.model.name

    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"更新AI配置失败：{str(e)}")

    return ai_config
# ----------------- 接口：删除模板（新增） -----------------
@router.delete("/templates/{template_id}")
async def delete_template(
    template_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    # 1. 校验模板是否存在且属于当前用户
    template = db.query(Template).filter(
        Template.id == template_id,
        Template.user_id == current_user.id,
        Template.status == "active"
    ).first()
    
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在或无权删除")

    # 2. 校验模板是否被公文使用（避免删除已关联的模板）
    used_by_doc = db.query(DocumentHistory).filter(
        DocumentHistory.template_id == template_id,
        DocumentHistory.user_id == current_user.id
    ).first()
    
    if used_by_doc:
        raise HTTPException(
            status_code=400,
            detail=f"该模板已被公文「{used_by_doc.filename}」使用，无法删除"
        )

    # 3. 执行删除（逻辑删除，而非物理删除，便于后续恢复）
    try:
        template.status = "deleted"  # 逻辑删除（推荐），若需物理删除则用 db.delete(template)
        db.commit()
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"删除模板失败：{str(e)}")

    return JSONResponse({"message": f"成功删除模板「{template.original_name}」"})
# ----------------- 接口：测试文件写入权限（保留原功能） -----------------
@router.get("/test-write")
async def test_write():
    try:
        test_file = os.path.join(DOWNLOAD_DIR, "test_write.txt")
        with open(test_file, "w", encoding="utf-8") as f:
            f.write(f"测试写入权限 - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        return {"status": "success", "message": "generated_docs目录具备写入权限", "test_file": test_file}
    except Exception as e:
        return {"status": "error", "message": f"写入权限测试失败：{str(e)}"}