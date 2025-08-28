from docx import Document
import os
import uuid

def save_uploaded_file(upload_file, upload_dir: str) -> str:
    ext = os.path.splitext(upload_file.filename)[1]
    filename = f"{uuid.uuid4().hex}{ext}"
    path = os.path.join(upload_dir, filename)
    with open(path, 'wb') as f:
        content = upload_file.file.read()
        f.write(content)
    return path

def render_docx_from_template(template_path: str | None, content: str) -> str:
    doc = Document()
    if template_path and os.path.exists(template_path):
        try:
            doc = Document(template_path)
        except Exception:
            doc = Document()
    doc.add_paragraph(content)
    out_name = f"generated_{uuid.uuid4().hex}.docx"
    out_path = os.path.join(os.path.dirname(__file__), '..', 'uploads', out_name)
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    return out_path
