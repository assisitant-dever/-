# utils.py
import json
import os
from docx import Document
import uuid
from .models import AIModelResponse,AIModel
DATA_FILE = os.path.join(os.path.dirname(__file__), "conversations.json")

async def save_uploaded_file(upload_file, upload_dir: str) -> str:
    ext = os.path.splitext(upload_file.filename)[1]
    filename = f"{uuid.uuid4().hex}{ext}"
    path = os.path.join(upload_dir, filename)
    with open(path, 'wb') as f:
        content = upload_file.file.read()
        f.write(content)
    return path

def render_docx_from_template(template_path: str | None, content: str) -> str:
    doc = Document()
    if template_path and os.path.exists(template_path):
        try:
            doc = Document(template_path)
        except Exception:
            doc = Document()
    doc.add_paragraph(content)
    out_name = f"generated_{uuid.uuid4().hex}.docx"
    out_path = os.path.join(os.path.dirname(__file__), '..', 'uploads', out_name)
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    return out_path

def load_data():
    if not os.path.exists(DATA_FILE):
        return []
    with open(DATA_FILE, "r", encoding="utf-8") as f:
        return json.load(f)

def save_data(data):
    with open(DATA_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def ai_model_to_response(ai_model: "AIModel", system_model: Optional["Model"] = None) -> AIModelResponse:
    """
    将 SQLAlchemy 的 AIModel 对象转换为 Pydantic 的 AIModelResponse。
    
    参数:
        ai_model: AIModel 实例
        system_model: 关联的系统模型 Model 实例（可选，如果没传会尝试从 ai_model.model 获取）
    
    返回:
        AIModelResponse 对象
    """
    # 如果没有传 system_model，则尝试从 ai_model.model 获取
    if not system_model:
        system_model = getattr(ai_model, "model", None)

    platform_name = "未知平台"
    model_name = "未知模型"

    if system_model:
        model_name = getattr(system_model, "name", "未知模型")
        platform = getattr(system_model, "platform", None)
        if platform:
            platform_name = getattr(platform, "name", "未知平台")

    # 处理脱敏 api_key
    raw_key = getattr(ai_model, "api_key", "") or ""
    if not raw_key:
        api_key_mask = ""
    elif len(raw_key) <= 8:
        api_key_mask = f"{raw_key[:4]}***"
    else:
        api_key_mask = f"{raw_key[:4]}***{raw_key[-4:]}"

    return AIModelResponse(
        id=ai_model.id,
        api_key_mask=api_key_mask,
        base_url=getattr(ai_model, "base_url", None),
        created_at=getattr(ai_model, "created_at", datetime.utcnow()),
        platform_name=platform_name,
        model_name=model_name
    )