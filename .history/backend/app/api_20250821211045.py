from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Depends
from fastapi.responses import FileResponse, JSONResponse
import os, time
from docx import Document
from dotenv import load_dotenv
from sqlalchemy.orm import Session

from .qwen_client import generate_text
from .utils import save_uploaded_file
from .database import SessionLocal
from .models import DocumentHistory, Template
from .deps import get_current_user

load_dotenv()

router = APIRouter()

# 上传模板目录
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), '..', 'uploads')
os.makedirs(UPLOAD_DIR, exist_ok=True)

# 生成文件保存目录
DOWNLOAD_DIR = os.path.join(os.getcwd(), "generated_docs")
os.makedirs(DOWNLOAD_DIR, exist_ok=True)

# ----------------- PROMPTS -----------------
PROMPTS = {
    "通知": "你是一个能写正式公文的助手，按照中国公文格式写一份{type}，注意语言正式、结构清晰，包含时间、地点、签发人等必要要素。",
    "请示": "你是一个能写正式公文的助手，写一份{type}，开门见山说明请求事由、请求事项、依据与建议。",
    "会议纪要": "写一份{type}，包含参会人员、时间地点、讨论要点、决议与后续事项。",
}

# ----------------- 数据库 Session -----------------
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# ----------------- 上传模板 -----------------
@router.post("/upload-template")
async def upload_template(
    file: UploadFile = File(...),
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    path = await save_uploaded_file(file, UPLOAD_DIR)
    
    # 保存模板记录到数据库
    template = Template(
        user_id=current_user.id,
        filename=os.path.basename(path),
        original_name=file.filename
    )
    db.add(template)
    db.commit()
    db.refresh(template)
    
    return {"filename": template.filename, "original_name": template.original_name}

# ----------------- 生成公文 -----------------
@router.post("/generate")
async def generate_document(
    doc_type: str = Form(...),
    user_input: str = Form(...),
    template_filename: str | None = Form(None),
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    base_prompt = PROMPTS.get(doc_type, "请写一份正式公文：{type}")
    prompt = base_prompt.format(type=doc_type) + "\n用户要求：" + user_input

    # 调用 AI 生成文本
    generated = generate_text(prompt)

    # 生成 DOCX 文件
    doc = Document()
    for line in generated.split("\n"):
        line = line.strip()
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        else:
            p = doc.add_paragraph()
            parts = line.split("**")
            for i, part in enumerate(parts):
                if i % 2 == 1:
                    run = p.add_run(part)
                    run.bold = True
                else:
                    subparts = part.split("*")
                    for j, sp in enumerate(subparts):
                        run = p.add_run(sp)
                        if j % 2 == 1:
                            run.italic = True

    filename = f"{doc_type}_{int(time.time())}.docx"
    file_path = os.path.join(DOWNLOAD_DIR, filename)
    doc.save(file_path)

    # 保存到数据库
    doc_record = DocumentHistory(
        user_id=current_user.id,
        doc_type=doc_type,
        content=generated,
        filename=filename
    )
    db.add(doc_record)
    db.commit()
    db.refresh(doc_record)

    # 返回文本和文件名
    return JSONResponse({"text": generated, "filename": filename})

# ----------------- 下载 DOCX -----------------
@router.get("/download/{filename}")
async def download(filename: str, current_user = Depends(get_current_user)):
    file_path = os.path.join(DOWNLOAD_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件未找到")
    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )

# ----------------- 获取用户文档历史 -----------------
@router.get("/history")
async def get_history(current_user = Depends(get_current_user), db: Session = Depends(get_db)):
    docs = db.query(DocumentHistory).filter(DocumentHistory.user_id == current_user.id).order_by(DocumentHistory.created_at.desc()).all()
    return [{"id": d.id, "doc_type": d.doc_type, "filename": d.filename, "created_at": d.created_at} for d in docs]

# ----------------- 获取用户模板列表 -----------------
@router.get("/templates")
async def get_templates(current_user = Depends(get_current_user), db: Session = Depends(get_db)):
    templates = db.query(Template).filter(Template.user_id == current_user.id).order_by(Template.uploaded_at.desc()).all()
    return [{"id": t.id, "filename": t.filename, "original_name": t.original_name, "uploaded_at": t.uploaded_at} for t in templates]


@router.get("/documents")
async def list_documents():
    files = os.listdir(DOWNLOAD_DIR)
    files = [f for f in files if f.endswith(".docx")]
    return {"documents": files}