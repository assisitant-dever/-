# api.py
from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Depends, Query
from fastapi.responses import FileResponse, JSONResponse,StreamingResponse
from fastapi.concurrency import run_in_threadpool
import json
import os, time
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
from sqlalchemy.orm import Session, joinedload
from sqlalchemy.future import select
from sqlalchemy.ext.asyncio import AsyncSession
import asyncio
from pydantic import ValidationError
from typing import List, Optional
from datetime import datetime
import pytz
import traceback
import urllib
import re
# 导入自定义模块（确保路径正确）
from .utils import save_uploaded_file
from .AI_client import generate_text_for_user, get_user_model_preference  # 多平台Client核心函数
from .database import SessionLocal, get_db,get_async_db 
from .models import (
    DocumentHistory, Template, AIModel, Conversation, Message, Platform,AIModelUpdate,
    AIModelResponse, PlatformModelResponse, User, AIModelCreate, SystemModelResponse, Model
)
from .deps import get_current_user
from .encryption import encrypt_api_key, decrypt_api_key
import logging
from fastapi.responses import JSONResponse
from fastapi.exceptions import RequestValidationError
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[logging.FileHandler("api.log"), logging.StreamHandler()]
)
logger = logging.getLogger(__name__)
# 加载环境变量
load_dotenv()

router = APIRouter()

# ----------------- 目录配置（保持原路径逻辑） -----------------
# 上传模板目录
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), '..', 'uploads')
os.makedirs(UPLOAD_DIR, exist_ok=True)

# 生成文件保存目录
DOWNLOAD_DIR = os.path.join(os.getcwd(), "generated_docs")
os.makedirs(DOWNLOAD_DIR, exist_ok=True)

# ----------------- 公文生成Prompt配置（保持原逻辑） -----------------
PROMPTS = {
    "通知": "你是专业的正式公文写作助手，需严格按照《党政机关公文处理工作条例》规定的通知格式撰写。内容需包含标题（发文机关+事由+文种）、主送机关、正文（明确通知缘由、具体事项、执行要求，涉及任务需标注责任主体）、附件说明（如有附件需注明附件名称及数量）、发文机关署名、成文日期、签发人（上行文需标注），语言需庄重严谨、逻辑清晰，避免歧义。",
    "请示": "你是专业的正式公文写作助手，需严格遵循请示“一文一事、单头主送”的核心要求撰写。内容需包含标题（发文机关+事由+文种）、主送机关（仅限一个直接上级机关）、正文（开门见山说明请示缘由，需引用政策依据或实际需求；清晰列出请示事项，做到具体明确；提出合理可行的建议方案供上级决策）、发文机关署名、成文日期、签发人，语言需谦恭得体，避免多头请示或夹带报告事项。",
    "会议纪要": "你是专业的正式公文写作助手，需按规范会议纪要格式撰写，确保内容客观准确、要素完整。需包含标题（会议名称+文种，如“XX会议纪要”）、基本信息（会议时间、会议地点、主持人、记录人、参会人员/列席人员/缺席人员名单及职务）、正文（分点梳理会议讨论要点，需准确反映不同观点；明确会议形成的决议事项，标注每项决议的责任单位、完成时限）、结尾（如需补充说明未尽事宜或后续沟通方式，可简要标注），语言需简洁凝练，避免主观评价。",
    "报告": "你是专业的正式公文写作助手，需严格按照报告的法定格式与写作要求撰写，区分“汇报工作”“反映情况”“答复询问”等不同报告类型。内容需包含标题（发文机关+事由+文种，如“XX关于XX工作的报告”）、主送机关、正文（开头简要说明报告背景或依据，主体部分详细阐述工作进展、成果、存在问题、原因分析及下一步计划，答复类报告需直接回应上级询问事项）、发文机关署名、成文日期、签发人（上行文需标注），语言需客观真实，不得夹带请示事项。",
    "函": "你是专业的正式公文写作助手，需根据函的“不相隶属机关商洽工作、询问答复问题、请求批准事项”功能，按规范格式撰写。内容需包含标题（发文机关+事由+文种，如“XX关于XX事宜的函”或“XX关于XX问题的复函”）、主送机关（明确的不相隶属机关）、正文（商洽类需说明事由及具体需求，询问类需列出清晰的问题清单，答复类需直接回应来函事项，请求批准类需说明依据及具体请求）、发文机关署名、成文日期，语言需平和得体，避免命令性或指令性表述。",
    "对话": "你是专业的公文场景对话撰写助手，需围绕公文办理全流程（如公文起草沟通、审批意见反馈、事项协调对接等）撰写对话内容。需明确对话场景（如“公文起草小组沟通会议对话”“上下级机关审批意见反馈对话”）、对话主体（标注角色及职务，如“起草人-XX部科员”“审批人-XX局副局长”）、对话逻辑（需符合公文办理规范，内容需聚焦具体事项，如格式修改建议、内容补充要求、办理时限确认等），语言需贴合职场沟通语境，既保持正式性，又体现沟通的针对性与高效性。"
}

# ----------------- 核心接口：上传模板（增强安全校验） -----------------
@router.post("/upload-template")
async def upload_template(
    file: UploadFile = File(...),
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # 新增1：限制文件大小（例如最大10MB）
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    file_size = 0
    try:
        # 读取文件前先检查大小
        while chunk := await file.read(1024 * 1024):  # 1MB块读取
            file_size += len(chunk)
            if file_size > MAX_FILE_SIZE:
                raise HTTPException(
                    status_code=413, 
                    detail=f"文件过大，最大支持{MAX_FILE_SIZE//1024//1024}MB"
                )
        # 重置文件指针（否则后续保存会是空文件）
        await file.seek(0)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"文件读取失败：{str(e)}")
    
    # 原有格式校验增强：同时检查MIME类型和文件头
    if not (file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" 
            and file.filename and file.filename.endswith(".docx")):
        raise HTTPException(status_code=400, detail="请上传正确的docx格式文件")
    
    # 保存文件
    file_path = await save_uploaded_file(file, UPLOAD_DIR)
    filename = os.path.basename(file_path)
    
    # 新增2：验证docx文件合法性（避免伪装成docx的恶意文件）
    try:
        # 尝试用python-docx打开，验证文件结构
        doc = Document(file_path)
        template_content = "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        # 验证失败则删除无效文件
        if os.path.exists(file_path):
            os.remove(file_path)
        raise HTTPException(status_code=400, detail=f"无效的docx文件：{str(e)}")
    
    # 保存模板记录
    template = Template(
        user_id=current_user.id,
        filename=filename,
        original_name=file.filename,
        content=template_content,
        status="active"
    )
    db.add(template)
    db.commit()
    db.refresh(template)
    
    return {
        "id": template.id,
        "filename": template.filename,
        "original_name": template.original_name,
        "uploaded_at": template.uploaded_at
    }

# ----------------- 核心接口：公文生成（适配多平台Client+数据库会话） -----------------
@router.post("/generate")
async def generate_document(
    doc_type: str = Form(...),
    user_input: str = Form(...),
    conv_id: Optional[int] = Form(None),
    ai_model_id: Optional[int] = Form(None),
    template_id: Optional[int] = Form(None),
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    公文生成接口（默认流式输出）
    - 流式返回AI生成的文本片段（SSE格式）
    - 生成完成后返回文件、会话等元数据
    """
    # -------------------------- 1. 前置校验与Prompt组装 --------------------------
    # 1.1 组装公文Prompt（含模板内容）
    base_prompt = PROMPTS.get(doc_type, f"请写一份正式公文：{doc_type}")
    if template_id:
        # 校验模板归属与有效性
        template = db.query(Template).filter(
            Template.id == template_id,
            Template.user_id == current_user.id,
            Template.status == "active"
        ).first()
        if not template:
            raise HTTPException(status_code=404, detail="指定模板不存在或无权访问")
        prompt = f"{base_prompt}\n模板内容：{template.content}\n用户要求：{user_input}"
    else:
        prompt = f"{base_prompt}\n用户要求：{user_input}"

    # 1.2 获取用户AI模型配置（优先手动选择，其次默认偏好）
    selected_ai_model: Optional[AIModel] = None
    try:
        if ai_model_id:
            # 优先使用用户手动选择的模型（预加载Platform关联）
            selected_ai_model = db.query(AIModel).filter(
                AIModel.id == ai_model_id,
                AIModel.user_id == current_user.id
            ).options(
                joinedload(AIModel.model).joinedload(Model.platform)
            ).first()
            if not selected_ai_model:
                raise HTTPException(status_code=404, detail="选择的AI模型不存在或无权访问")

            # 平台名称映射（数据库存储名 → 客户端工厂标识）
            platform_mapping = {
                "Alibaba": "qwen",
                "OpenAI": "openai",
                "Anthropic": "anthropic",
                "Google": "gemini"
            }
            db_platform = selected_ai_model.model.platform.name
            user_config = {
                "provider": platform_mapping.get(db_platform, db_platform.lower()),
                "api_key": decrypt_api_key(selected_ai_model.api_key),
                "model_name": selected_ai_model.model.name,
                "base_url": selected_ai_model.effective_base_url  # 假设models中定义了effective_base_url属性
            }

        else:
            # 无手动选择时，使用用户默认偏好（上次使用/最新添加）
            user_config = get_user_model_preference( 
                user_id=current_user.id, db=db
            )
            if not user_config:
                raise HTTPException(status_code=400, detail="请先在「API Keys管理」中添加AI模型配置")

            # 平台名称映射（同手动选择逻辑）
            platform_mapping = {
                "Alibaba": "qwen",
                "OpenAI": "openai",
                "Anthropic": "anthropic",
                "Google": "gemini"
            }
            user_config["provider"] = platform_mapping.get(
                user_config["provider"],
                user_config["provider"].lower()
            )

            # 同步获取模型实例（用于后续会话关联）
            selected_ai_model = db.query(AIModel).filter(
                AIModel.model.has(
                    Model.platform.has(name=user_config["provider"]),
                    name=user_config["model_name"]
                ),
                AIModel.user_id == current_user.id
            ).options(
                joinedload(AIModel.model).joinedload(Model.platform)
            ).first()
            if not selected_ai_model:
                raise HTTPException(status_code=500, detail="用户默认模型配置异常，请重新添加")

    except HTTPException as e:
        raise e  # 直接抛出前置校验异常
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取模型配置失败：{str(e)}")

    # -------------------------- 2. 初始化流式生成 --------------------------
    # 用线程池包装同步函数（避免阻塞FastAPI事件循环）
    try:
        generated_iterator: Iterator[str] = await run_in_threadpool(
            generate_text_for_user,
            user_id=current_user.id,
            prompt=prompt,
            system_prompt=base_prompt,
            db=db,
            stream=True  # 强制开启流式（默认行为）
        )
    except Exception as e:
        error_msg = str(e)
        # 格式化常见错误
        if "401" in error_msg and "invalid_api_key" in error_msg.lower():
            raise HTTPException(status_code=401, detail="API Key不正确或已失效")
        elif "不支持的平台" in error_msg:
            raise HTTPException(status_code=400, detail="所选AI平台暂不支持")
        else:
            raise HTTPException(status_code=500, detail=f"流式生成初始化失败：{error_msg}")

    # -------------------------- 3. 定义SSE流式生成器 --------------------------
    full_content: list[str] = []  # 收集完整内容（用于后续DOCX生成和数据库存储）

    async def sse_generator():
        nonlocal full_content
        try:
            # 3.1 实时返回流式文本片段
            for chunk in generated_iterator:
                if chunk.strip():  # 过滤空片段
                    full_content.append(chunk)
                    # 按SSE规范返回（data字段+JSON序列化，避免前端解析异常）
                    yield f"data: {json.dumps({'chunk': chunk})}\n\n"
                    # 可选：控制流速度（避免前端处理过快）
                    await asyncio.sleep(0.01)

            # 3.2 流式结束后，处理完整内容（DOCX生成+数据库存储）
            generated_full = "".join(full_content)
            if not generated_full:
                yield f"event: error\ndata: {json.dumps({'detail': 'AI生成内容为空'})}\n\n"
                return

            # -------------------------- 3.2.1 生成DOCX文件 --------------------------
            doc = Document()
            # 设置公文标准格式（仿宋GB2312、小四号字、无段间距）
            normal_style = doc.styles["Normal"]
            font = normal_style.font
            font.name = "FangSong_GB2312"
            font.size = Pt(16)
            normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "FangSong_GB2312")
            normal_style.paragraph_format.space_before = Pt(0)
            normal_style.paragraph_format.space_after = Pt(0)

            # 解析Markdown格式（标题、加粗、斜体）
            for line in generated_full.split("\n"):
                line = line.strip()
                if not line:
                    continue
                # 处理标题（# 一级 / ## 二级 / ### 三级）
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                else:
                    # 处理加粗（**内容**）和斜体（*内容*）
                    para = doc.add_paragraph()
                    bold_parts = line.split("**")
                    for i, part in enumerate(bold_parts):
                        if i % 2 == 1:  # 奇数段为加粗内容
                            run = para.add_run(part)
                            run.bold = True
                        else:
                            italic_parts = part.split("*")
                            for j, sub_part in enumerate(italic_parts):
                                run = para.add_run(sub_part)
                                if j % 2 == 1:  # 奇数段为斜体内容
                                    run.italic = True

            # 保存DOCX（用户ID+时间戳避免冲突）
            os.makedirs(DOWNLOAD_DIR, exist_ok=True)  # 确保目录存在
            filename = f"{doc_type}_{current_user.id}_{int(time.time())}.docx"
            file_path = os.path.join(DOWNLOAD_DIR, filename)
            doc.save(file_path)

            # -------------------------- 3.2.2 保存数据库记录 --------------------------
            # 1. 保存公文历史
            doc_record = DocumentHistory(
                user_id=current_user.id,
                doc_type=doc_type,
                content=generated_full,
                filename=filename,
                template_id=template_id
            )
            db.add(doc_record)
            db.flush()  # 提前获取doc_id，避免依赖commit

            # 2. 处理会话（创建/更新）
            if conv_id:
                # 更新已有会话
                conversation = db.query(Conversation).filter(
                    Conversation.id == conv_id,
                    Conversation.user_id == current_user.id
                ).first()
                if not conversation:
                    yield f"event: error\ndata: {json.dumps({'detail': '指定会话不存在'})}\n\n"
                    db.rollback()
                    return
                conversation.updated_at = datetime.now(pytz.UTC)
                conversation.ai_model_id = selected_ai_model.id
                conversation.status = "active"
            else:
                # 创建新会话（生成简短标题）
                input_short = user_input.strip()[:8] if len(user_input.strip()) > 8 else user_input.strip()
                conversation = Conversation(
                    user_id=current_user.id,
                    ai_model_id=selected_ai_model.id,
                    title=f"{doc_type}生成_{input_short}...",
                    status="active",
                    created_at=datetime.now(pytz.UTC),
                    updated_at=datetime.now(pytz.UTC)
                )
                db.add(conversation)
            db.flush()  # 提前获取conversation_id

            # 3. 保存消息记录（用户输入+AI回复）
            user_msg = Message(
                conversation_id=conversation.id,
                role="user",
                content=user_input,
                created_at=datetime.now(pytz.UTC)
            )
            ai_msg = Message(
                conversation_id=conversation.id,
                role="assistant",
                content=generated_full,
                docx_file=filename,
                created_at=datetime.now(pytz.UTC)
            )
            db.add_all([user_msg, ai_msg])

            # 提交所有数据库操作
            db.commit()
            db.refresh(doc_record)
            db.refresh(conversation)

            # -------------------------- 3.2.3 发送元数据事件 --------------------------
            # 包含文件下载、会话续接所需信息
            metadata = {
                "filename": filename,
                "conv_id": conversation.id,
                "doc_id": doc_record.id,
                "used_model": f"{selected_ai_model.model.platform.name} - {selected_ai_model.model.name}",
                "full_text": generated_full  
            }
            yield f"event: metadata\ndata: {json.dumps(metadata)}\n\n"
            await asyncio.sleep(0.001)  # 强制推送

            # -------------------------- 3.2.4 发送生成完成事件 --------------------------
            yield f"event: complete\ndata: {json.dumps({'status': 'success', 'msg': '生成完成'})}\n\n"
            await asyncio.sleep(0.001)
        except Exception as e:
            # 捕获生成过程中的异常，发送错误事件并回滚
            error_detail = str(e)
            # 格式化常见错误
            if "permission denied" in error_detail.lower():
                error_detail = "文件保存失败，请检查服务器权限"
            elif "sqlalchemy" in error_detail.lower():
                error_detail = "数据库操作失败，请稍后重试"
            else:
                error_detail = f"生成异常：{error_detail}"

            # 发送错误事件
            yield f"event: error\ndata: {json.dumps({'detail': error_detail})}\n\n"
            # 回滚未提交的数据库操作
            db.rollback()
        finally:
            # 关闭数据库会话（根据依赖注入逻辑调整，避免连接泄漏）
            db.close()

    # -------------------------- 4. 返回StreamingResponse --------------------------
    return StreamingResponse(
        sse_generator(),
        media_type="text/event-stream",  # SSE标准媒体类型
        headers={
            "Cache-Control": "no-cache",  # 禁止客户端缓存
            "Connection": "keep-alive",  # 保持长连接
            "X-Accel-Buffering": "no"  # 禁用Nginx等反向代理的缓冲（关键！确保实时性）
        }
    )

# ----------------- 接口：下载生成的DOCX文件（增强安全） -----------------
@router.get("/download/{filename}")
async def download(
    filename: str, 
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    decoded_filename = urllib.parse.unquote(filename)
    
    # 新增：严格校验文件名格式，防止路径遍历（如../../etc/passwd）
    if not re.match(r'^[\w\-]+\.docx$', decoded_filename):  # 只允许字母、数字、下划线、连字符和.docx后缀
        raise HTTPException(status_code=400, detail="文件名格式非法")
    
    file_path = os.path.join(DOWNLOAD_DIR, decoded_filename)
    
    # 新增：使用realpath确保文件在DOWNLOAD_DIR内
    if not os.path.realpath(file_path).startswith(os.path.realpath(DOWNLOAD_DIR)):
        raise HTTPException(status_code=403, detail="访问受限")
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件未找到或已被删除")
    
    # 数据库校验
    doc_record = db.query(DocumentHistory).filter(
        DocumentHistory.filename == decoded_filename,
        DocumentHistory.user_id == current_user.id
    ).first()
    if not doc_record:
        raise HTTPException(status_code=403, detail="无权访问此文件")
    
    encoded_filename = urllib.parse.quote(decoded_filename)
    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )

    

# ----------------- 接口：获取用户公文历史 -----------------
# ----------------- 接口：获取用户公文历史（新增分页） -----------------
@router.get("/history")
async def get_history(
    current_user = Depends(get_current_user), 
    db: Session = Depends(get_db),
    page: int = Query(1, ge=1, description="页码，从1开始"),  # 新增分页参数
    page_size: int = Query(10, ge=1, le=50, description="每页条数，最大50")  # 限制最大条数
):
    # 计算偏移量
    offset = (page - 1) * page_size
    
    # 1. 先查询总数（用于前端分页控件）
    total = db.query(DocumentHistory).filter(
        DocumentHistory.user_id == current_user.id
    ).count()
    
    # 2. 分页查询数据（仅加载当前页数据）
    docs = db.query(DocumentHistory).filter(
        DocumentHistory.user_id == current_user.id
    ).order_by(DocumentHistory.created_at.desc())\
     .offset(offset).limit(page_size).all()  # 分页核心逻辑
    
    # 补充关联的模板名称
    result = []
    for doc in docs:
        template_name = None
        if doc.template_id:
            # 优化：使用exists查询替代完整查询，减少IO
            template_exists = db.query(Template).filter(
                Template.id == doc.template_id,
                Template.status == "active"
            ).exists()
            if db.query(template_exists).scalar():
                template = db.query(Template).filter(
                    Template.id == doc.template_id
                ).with_entities(Template.original_name).first()
                template_name = template.original_name if template else None
        
        result.append({
            "id": doc.id,
            "doc_type": doc.doc_type,
            "filename": doc.filename,
            "used_template": template_name,
            "created_at": doc.created_at.strftime("%Y-%m-%d %H:%M:%S"),
            "content_preview": doc.content[:50] + "..." if len(doc.content) > 50 else doc.content
        })
    
    # 返回分页元数据+数据列表
    return {
        "total": total,
        "page": page,
        "page_size": page_size,
        "data": result
    }


# ----------------- 接口：获取用户模板列表（新增分页） -----------------
@router.get("/templates")
async def get_templates(
    current_user = Depends(get_current_user), 
    db: Session = Depends(get_db),
    page: int = Query(1, ge=1),
    page_size: int = Query(10, ge=1, le=50)
):
    offset = (page - 1) * page_size
    total = db.query(Template).filter(
        Template.user_id == current_user.id,
        Template.status == "active"
    ).count()
    
    templates = db.query(Template).filter(
        Template.user_id == current_user.id,
        Template.status == "active"
    ).order_by(Template.uploaded_at.desc())\
     .offset(offset).limit(page_size).all()
    
    return {
        "total": total,
        "page": page,
        "page_size": page_size,
        "data": [
            {
                "id": t.id,
                "filename": t.filename,
                "original_name": t.original_name,
                "uploaded_at": t.uploaded_at.strftime("%Y-%m-%d %H:%M:%S"),
                "content_preview": t.content[:30] + "..." if len(t.content) > 30 else t.content
            } for t in templates
        ]
    }

# ----------------- 接口：获取用户模板列表 -----------------
@router.get("/templates")
async def get_templates(
    current_user = Depends(get_current_user), 
    db: Session = Depends(get_db)
):
    # 查询当前用户的有效模板（按上传时间倒序）
    templates = db.query(Template).filter(
        Template.user_id == current_user.id,
        Template.status == "active"  # 只返回有效模板
    ).order_by(Template.uploaded_at.desc()).all()
    
    return [
        {
            "id": t.id,
            "filename": t.filename,
            "original_name": t.original_name,
            "uploaded_at": t.uploaded_at.strftime("%Y-%m-%d %H:%M:%S"),
            "content_preview": t.content[:30] + "..." if len(t.content) > 30 else t.content  # 模板内容预览
        }
        for t in templates
    ]

# ----------------- 接口：获取模板内容（从数据库读取，替代文件读取） -----------------
@router.get("/template-content/{template_id}")  # 改用template_id，避免文件名依赖
async def get_template_content(
    template_id: int, 
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # 从数据库查询有效模板（含内容），避免重复读取文件
    template = db.query(Template).filter(
        Template.id == template_id,
        Template.user_id == current_user.id,
        Template.status == "active"
    ).first()
    
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在或无权访问")
    
    # 直接返回数据库中的内容（无需再读取文件）
    return JSONResponse({
        "id": template.id,
        "original_name": template.original_name,
        "content": template.content
    })

# ----------------- 接口：获取用户API Key列表（解密脱敏） -----------------
@router.get("/keys", response_model=List[AIModelResponse])
async def get_user_ai_configs(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_async_db)
):
    try:
        # 1. 查询并预加载关联
        stmt = select(AIModel).filter(
            AIModel.user_id == current_user.id
        ).options(
            joinedload(AIModel.model).joinedload(Model.platform)
        ).order_by(AIModel.created_at.desc())
        
        result = await db.execute(stmt)
        user_configs = result.scalars().unique().all()

        # 2. 手动提取字段值并构造验证数据（关键步骤）
        validated_data = []
        for config in user_configs:
            # 手动提取 model.name（确保非空）
            model_name = "未知模型"
            if config.model and hasattr(config.model, "name"):
                model_name = config.model.name or "未知模型"  # 处理空字符串
            
            # 手动提取 model.platform.name（确保非空）
            platform_name = "未知平台"
            if config.model and config.model.platform and hasattr(config.model.platform, "name"):
                platform_name = config.model.platform.name or "未知平台"  # 处理空字符串

            # 脱敏 API Key
            try:
                decrypted_key = decrypt_api_key(config.api_key) if config.api_key else ""
            except:
                decrypted_key = ""
            if len(decrypted_key) <= 8 and decrypted_key:
                api_key_mask = f"{decrypted_key[:4]}***"
            elif decrypted_key:
                api_key_mask = f"{decrypted_key[:4]}***{decrypted_key[-4:]}"
            else:
                api_key_mask = "***"

            # 构造符合 AIModelResponse 结构的字典
            validated_data.append({
                "id": config.id,
                "platform_name": platform_name,  # 使用手动提取的值
                "model_name": model_name,        # 使用手动提取的值
                "api_key_mask": api_key_mask,
                "base_url": config.base_url,
                "created_at": config.created_at
            })

        # 3. 用 Pydantic 验证构造的数据（此时字段均已存在）
        return [AIModelResponse(**data) for data in validated_data]

    except ValidationError as e:
        print(f"响应模型验证失败：{e}")
        raise HTTPException(status_code=500, detail="数据格式错误")
    except Exception as e:
        print(f"接口错误：{str(e)}")
        raise HTTPException(status_code=500, detail="获取配置失败")
    
# ----------------- 接口：获取支持的AI平台及模型列表 -----------------
@router.get("/platforms", response_model=List[PlatformModelResponse])
async def get_system_platforms(
    include_details: bool = Query(False, description="是否返回模型详情（如描述、是否支持）"),
    db: Session = Depends(get_db)
):
    """
    获取系统支持的AI平台列表（从数据库 Platform/Model 表查询，替代原硬编码）
    - include_details=True：返回模型详情（供前端选择模型时展示描述）
    - include_details=False：仅返回平台名称和模型名称列表（轻量查询）
    """
    # 只查询系统启用的平台（is_active=True），预加载关联的模型
    platforms = db.query(Platform).filter(Platform.is_active == True).options(
        joinedload(Platform.models)  # 预加载平台下的模型
    ).order_by(Platform.name.asc()).all()
    
    # 调用 Platform.to_response() 转换为 Pydantic 响应格式
    return [platform.to_response(include_details=include_details) for platform in platforms]


@router.get("/platforms/{platform_id}/models", response_model=List[SystemModelResponse])
async def get_platform_models(
    platform_id: int,
    db: Session = Depends(get_db)
):
    """
    获取指定平台下的所有系统模型（供前端“按平台筛选模型”）
    """
    # 1. 校验平台是否存在且启用
    platform = db.query(Platform).filter(
        Platform.id == platform_id,
        Platform.is_active == True
    ).first()
    if not platform:
        raise HTTPException(status_code=404, detail="系统未支持该平台或平台已禁用")
    
    # 2. 查询该平台下的所有支持模型（is_supported=True）
    models = db.query(Model).filter(
        Model.platform_id == platform_id,
        Model.is_supported == True
    ).order_by(Model.name.asc()).all()
    
    # 补充平台名称（SystemModelResponse需要platform_name字段）
    for model in models:
        model.platform_name = platform.name
    
    return models

# ----------------- 接口：新增AI模型配置（加密存储） -----------------
@router.post("/keys", response_model=AIModelResponse)
async def create_ai_model(
    ai_model_create: AIModelCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    用户新增AI模型配置（支持系统模型ID或自定义平台/模型名称）
    """
    if ai_model_create.model_id:
        # 系统模型逻辑
        system_model = db.query(Model).filter(
            Model.id == ai_model_create.model_id,
            Model.is_supported == True
        ).options(joinedload(Model.platform)).first()
        if not system_model:
            raise HTTPException(status_code=400, detail="系统未支持该模型，无法创建配置")
        platform_name = system_model.platform.name
        model_name = system_model.name
        base_url = ai_model_create.base_url or system_model.platform.base_url
        model_id = system_model.id
    else:
        # 自定义模型逻辑
        if not ai_model_create.platform_name or not ai_model_create.model_name:
            raise HTTPException(status_code=400, detail="自定义配置必须提供 platform_name 和 model_name")

        # 查找或创建平台
        platform = db.query(Platform).filter_by(name=ai_model_create.platform_name).first()
        if not platform:
            platform = Platform(name=ai_model_create.platform_name, base_url=ai_model_create.base_url or "")
            db.add(platform)
            db.commit()
            db.refresh(platform)

        # 查找或创建模型
        model = db.query(Model).filter_by(name=ai_model_create.model_name, platform_id=platform.id).first()
        if not model:
            model = Model(name=ai_model_create.model_name, platform_id=platform.id)
            db.add(model)
            db.commit()
            db.refresh(model)

        platform_name = platform.name
        model_name = model.name
        base_url = ai_model_create.base_url
        model_id = model.id

    # 检查重复配置
    existing_config = db.query(AIModel).join(AIModel.model).join(Model.platform).filter(
        AIModel.user_id == current_user.id,
        Model.name == model_name,
        Platform.name == platform_name
    ).first()
    if existing_config:
        raise HTTPException(
            status_code=400,
            detail=f"您已为「{platform_name}-{model_name}」配置过API Key，无需重复添加"
        )

    # 加密 API Key 并创建配置
    try:
        encrypted_key = encrypt_api_key(ai_model_create.api_key)
        new_ai_config = AIModel(
            user_id=current_user.id,
            model_id=model_id,
            api_key=encrypted_key,
            base_url=base_url
        )
        db.add(new_ai_config)
        db.commit()
        db.refresh(new_ai_config)

        # 脱敏 API Key
        decrypted_key = decrypt_api_key(new_ai_config.api_key)
        if len(decrypted_key) <= 8:
            new_ai_config.api_key_mask = f"{decrypted_key[:4]}***"
        else:
            new_ai_config.api_key_mask = f"{decrypted_key[:4]}***{decrypted_key[-4:]}"
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"创建AI配置失败：{str(e)}")

    # 返回响应（FastAPI + AIModelResponse 自动映射 property）
    return new_ai_config



# ----------------- 接口：删除AI模型配置 -----------------
@router.delete("/keys/{key_id}")
async def delete_ai_model(
    key_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    # 1. 查询待删除的模型（校验归属+预加载关联）
    ai_model = db.query(AIModel).filter(
        AIModel.id == key_id,
        AIModel.user_id == current_user.id
    ).options(
        joinedload(AIModel.model).joinedload(Model.platform)
    ).first()
    
    if not ai_model:
        raise HTTPException(status_code=404, detail="API配置不存在或无权删除")
    
    # 2. 校验是否有会话正在使用该模型（避免删除正在使用的模型）
    using_conversations = db.query(Conversation).filter(
        Conversation.user_id == current_user.id,
        Conversation.ai_model_id == key_id,
        Conversation.status == "active"
    ).first()
    
    if using_conversations:
        raise HTTPException(
            status_code=400, 
            detail=f"该模型正在被会话「{using_conversations.title}」使用，建议先修改会话的模型再删除"
        )

    # 3. 执行删除
    try:
        db.delete(ai_model)
        db.commit()
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail=f"删除AI模型配置失败：{str(e)}")

    # 返回完整的平台+模型名称（从关联表获取）
    platform_name = ai_model.model.platform.name
    model_name = ai_model.model.name
    return JSONResponse({"message": f"成功删除「{platform_name} - {model_name}」的API配置"})
# ----------------- 接口：编辑AI模型配置（新增PUT接口） -----------------
@router.put("/keys/{ai_model_id}", response_model=AIModelResponse)
async def update_ai_model(
    ai_model_id: int,
    ai_model_update: AIModelUpdate,  # Pydantic 模型
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    # 1️⃣ 检查 model_id 是否为空
    if not ai_model_update.model_id:
        raise HTTPException(
            status_code=400,
            detail="更新失败：被更新的model_id 不能为空"
        )

    # 2️⃣ 查询要更新的 AI 模型
    ai_model = db.query(AIModel).filter(
        AIModel.id == ai_model_id, AIModel.user_id == current_user.id
    ).first()
    if not ai_model:
        raise HTTPException(status_code=404, detail="AI 模型不存在")

    # 3️⃣ 查询系统模型是否存在且支持
    system_model = db.query(Model).filter(
        Model.id == ai_model_update.model_id,
        Model.is_supported == True
    ).first()
    if not system_model:
        raise HTTPException(status_code=400, detail="系统未支持该模型，无法切换")

    # 4️⃣ 更新字段
    ai_model.model_id = ai_model_update.model_id
    ai_model.api_key = ai_model_update.api_key
    if ai_model_update.base_url:
        ai_model.base_url = ai_model_update.base_url

    db.commit()
    db.refresh(ai_model)

    # 5️⃣ 构造返回字典
    response_data = {
        "id": ai_model.id,
        "api_key_mask": "",  # 默认脱敏为空
        "base_url": ai_model.base_url,
        "created_at": ai_model.created_at,
        "platform_name": system_model.platform.name if system_model.platform else "未知平台",
        "model_name": system_model.name or "未知模型"
    }

    # 6️⃣ 脱敏 API Key
    raw_key = ai_model.api_key or ""
    if raw_key:
        if len(raw_key) <= 8:
            response_data["api_key_mask"] = f"{raw_key[:4]}***"
        else:
            response_data["api_key_mask"] = f"{raw_key[:4]}***{raw_key[-4:]}"

    # 7️⃣ 返回 Pydantic 对象
    return AIModelResponse(**response_data)


# ----------------- 接口：更新模板内容（新增） -----------------
@router.put("/templates/{template_id}")
async def update_template(
    template_id: int,
    file: UploadFile = File(None),  # 可选：上传新文件更新内容
    new_name: Optional[str] = Form(None),  # 可选：仅修改名称
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    # 1. 校验模板存在性
    template = db.query(Template).filter(
        Template.id == template_id,
        Template.user_id == current_user.id,
        Template.status == "active"
    ).first()
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在或无权修改")
    
    # 2. 处理名称更新
    if new_name and new_name.strip():
        template.original_name = new_name.strip()
    
    # 3. 处理文件内容更新
    if file:
        # 复用上传时的安全校验逻辑
        MAX_FILE_SIZE = 10 * 1024 * 1024
        file_size = 0
        try:
            while chunk := await file.read(1024 * 1024):
                file_size += len(chunk)
                if file_size > MAX_FILE_SIZE:
                    raise HTTPException(status_code=413, detail="文件过大")
            await file.seek(0)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"文件读取失败：{str(e)}")
        
        if not (file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" 
                and file.filename and file.filename.endswith(".docx")):
            raise HTTPException(status_code=400, detail="请上传正确的docx格式文件")
        
        # 保存新文件并更新内容
        file_path = await save_uploaded_file(file, UPLOAD_DIR)
        template.filename = os.path.basename(file_path)
        try:
            doc = Document(file_path)
            template.content = "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            os.remove(file_path)
            raise HTTPException(status_code=400, detail=f"无效的docx文件：{str(e)}")
    
    # 4. 提交更新
    template.updated_at = datetime.now(pytz.UTC)  # 新增更新时间字段（需在models.Template中添加）
    db.commit()
    db.refresh(template)
    
    return {
        "id": template.id,
        "message": "模板更新成功",
        "original_name": template.original_name,
        "updated_at": template.updated_at.strftime("%Y-%m-%d %H:%M:%S")
    }
# ----------------- 接口：测试文件写入权限（保留原功能） -----------------
@router.get("/test-write")
async def test_write():
    try:
        test_file = os.path.join(DOWNLOAD_DIR, "test_write.txt")
        with open(test_file, "w", encoding="utf-8") as f:
            f.write(f"测试写入权限 - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        return {"status": "success", "message": "generated_docs目录具备写入权限", "test_file": test_file}
    except Exception as e:
        return {"status": "error", "message": f"写入权限测试失败：{str(e)}"}
@router.post("/conversations/{conversation_id}/generate_title")
async def generate_conversation_title(conversation_id: int, db: Session = Depends(get_db)):
    conversation = db.query(Conversation).filter(Conversation.id == conversation_id).first()
    if not conversation:
        raise HTTPException(status_code=404, detail="Conversation not found")

    messages = db.query(Message).filter(Message.conversation_id == conversation_id).order_by(Message.id).limit(2).all()
    content = " ".join([m.content for m in messages])
    print("生成标题用内容:", content)  # 🔍 打印看看

    try:
        title = await generate_text_for_user(content)  # 调用 AI
    except Exception as e:
        print("AI 生成标题出错:", e)
        raise HTTPException(status_code=500, detail="AI生成标题失败")

    conversation.title = title.strip()
    db.commit()
    print("新标题:", conversation.title)
    return {"id": conversation.id, "title": conversation.title}

