from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Depends
from fastapi.responses import FileResponse, JSONResponse
import os, time
from docx import Document
from dotenv import load_dotenv
from sqlalchemy.orm import Session
from .utils import load_data
from .qwen_client import generate_text
from .utils import save_uploaded_file
from .database import SessionLocal
from .models import DocumentHistory, Template
from .deps import get_current_user

load_dotenv()

router = APIRouter()

# ä¸Šä¼ æ¨¡æ¿ç›®å½•
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), '..', 'uploads')
os.makedirs(UPLOAD_DIR, exist_ok=True)
CONVERSATIONS_FILE = "conversations.json"

# ç”Ÿæˆæ–‡ä»¶ä¿å­˜ç›®å½•
DOWNLOAD_DIR = os.path.join(os.getcwd(), "generated_docs")
os.makedirs(DOWNLOAD_DIR, exist_ok=True)

# ----------------- PROMPTS -----------------
PROMPTS = {
    "é€šçŸ¥": "ä½ æ˜¯ä¸€ä¸ªèƒ½å†™æ­£å¼å…¬æ–‡çš„åŠ©æ‰‹ï¼ŒæŒ‰ç…§ä¸­å›½å…¬æ–‡æ ¼å¼å†™ä¸€ä»½{type}ï¼Œæ³¨æ„è¯­è¨€æ­£å¼ã€ç»“æ„æ¸…æ™°ï¼ŒåŒ…å«æ—¶é—´ã€åœ°ç‚¹ã€ç­¾å‘äººç­‰å¿…è¦è¦ç´ ã€‚",
    "è¯·ç¤º": "ä½ æ˜¯ä¸€ä¸ªèƒ½å†™æ­£å¼å…¬æ–‡çš„åŠ©æ‰‹ï¼Œå†™ä¸€ä»½{type}ï¼Œå¼€é—¨è§å±±è¯´æ˜è¯·æ±‚äº‹ç”±ã€è¯·æ±‚äº‹é¡¹ã€ä¾æ®ä¸å»ºè®®ã€‚",
    "ä¼šè®®çºªè¦": "å†™ä¸€ä»½{type}ï¼ŒåŒ…å«å‚ä¼šäººå‘˜ã€æ—¶é—´åœ°ç‚¹ã€è®¨è®ºè¦ç‚¹ã€å†³è®®ä¸åç»­äº‹é¡¹ã€‚",
}

# ----------------- æ•°æ®åº“ Session -----------------
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# ----------------- ä¸Šä¼ æ¨¡æ¿ -----------------
@router.post("/upload-template")
async def upload_template(
    file: UploadFile = File(...),
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    path = await save_uploaded_file(file, UPLOAD_DIR)
    
    # ä¿å­˜æ¨¡æ¿è®°å½•åˆ°æ•°æ®åº“
    template = Template(
        user_id=current_user.id,
        filename=os.path.basename(path),
        original_name=file.filename
    )
    db.add(template)
    db.commit()
    db.refresh(template)
    
    return {"filename": template.filename, "original_name": template.original_name}

@router.post("/generate")
async def generate_document(
    doc_type: str = Form(...),
    user_input: str = Form(...),
    conv_id: int = Form(...),  
    template_content: str = Form(None),
    template_filename: str | None = Form(None),
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    base_prompt = PROMPTS.get(doc_type, "è¯·å†™ä¸€ä»½æ­£å¼å…¬æ–‡ï¼š{type}")
    prompt = base_prompt.format(type=doc_type) + "\nç”¨æˆ·è¦æ±‚ï¼š" + user_input

    generated = generate_text(prompt)

    # ç”Ÿæˆ DOCX
    doc = Document()
    # ...ï¼ˆåŸæœ‰ä»£ç ï¼‰

    filename = f"{doc_type}_{int(time.time())}.docx"
    file_path = os.path.join(DOWNLOAD_DIR, filename)
    doc.save(file_path)

    # ä¿å­˜åˆ°æ•°æ®åº“
    doc_record = DocumentHistory(
        user_id=current_user.id,
        doc_type=doc_type,
        content=generated,
        filename=filename
    )
    db.add(doc_record)
    db.commit()

    # ğŸ” æŠŠç”Ÿæˆçš„å…¬æ–‡åŠ å…¥å¯¹è¯
    data = load_data()
    for conv in data:
        if conv["id"] == conv_id:
            msg = {
                "id": int(time.time() * 1000),
                "role": "assistant",
                "content": generated,
                "docx_file": filename,
                "created_at": time.strftime("%Y-%m-%d %H:%M:%S")
            }
            conv["messages"].append(msg)
            save_data(data)
            break

    return JSONResponse({"text": generated, "filename": filename},
                        status_code=200)

# ----------------- ä¸‹è½½ DOCX -----------------
@router.get("/download/{filename}")
async def download(filename: str, current_user = Depends(get_current_user)):
    file_path = os.path.join(DOWNLOAD_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="æ–‡ä»¶æœªæ‰¾åˆ°")
    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )

# ----------------- è·å–ç”¨æˆ·æ–‡æ¡£å†å² -----------------
@router.get("/history")
async def get_history(current_user = Depends(get_current_user), db: Session = Depends(get_db)):
    docs = db.query(DocumentHistory).filter(DocumentHistory.user_id == current_user.id).order_by(DocumentHistory.created_at.desc()).all()
    return [{"id": d.id, "doc_type": d.doc_type, "filename": d.filename, "created_at": d.created_at} for d in docs]

# ----------------- è·å–ç”¨æˆ·æ¨¡æ¿åˆ—è¡¨ -----------------
@router.get("/templates")
async def get_templates(current_user = Depends(get_current_user), db: Session = Depends(get_db)):
    templates = db.query(Template).filter(Template.user_id == current_user.id).order_by(Template.uploaded_at.desc()).all()
    return [{"id": t.id, "filename": t.filename, "original_name": t.original_name, "uploaded_at": t.uploaded_at} for t in templates]


@router.get("/documents")
async def list_documents():
    files = os.listdir(DOWNLOAD_DIR)
    files = [f for f in files if f.endswith(".docx")]
    return {"documents": files}

@router.get("/template-content/{filename}")
async def get_template_content(filename: str, current_user=Depends(get_current_user)):
    file_path = os.path.join(UPLOAD_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="æ¨¡æ¿æ–‡ä»¶æœªæ‰¾åˆ°")

    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return JSONResponse({"content": "\n".join(full_text)})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"è¯»å–æ¨¡æ¿å¤±è´¥: {str(e)}")