from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Depends
from fastapi.responses import FileResponse, JSONResponse
import os, time
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
from sqlalchemy.orm import Session
from .utils import load_data,save_data
from .qwen_client import generate_text
from .utils import save_uploaded_file
from .database import SessionLocal,get_db
from .models import DocumentHistory, Template
from .deps import get_current_user
import json
from fastapi import Request
from .conversations import _add_message_to_conversation, load_data, save_data
from typing import List
from .models import AIModelResponse,PlatformModelResponse,User,AIModelCreate,AIModel
from datetime import datetime,timezone
import pytz
from .encryption import encrypt_api_key,decrypt_api_key
load_dotenv()

router = APIRouter()

# 上传模板目录
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), '..', 'uploads')
os.makedirs(UPLOAD_DIR, exist_ok=True)

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CONVERSATIONS_FILE = os.path.join(PROJECT_ROOT, "conversations.json")

# 生成文件保存目录
DOWNLOAD_DIR = os.path.join(os.getcwd(), "generated_docs")
os.makedirs(DOWNLOAD_DIR, exist_ok=True)

# ----------------- PROMPTS -----------------
PROMPTS = {
    "通知": "你是专业的正式公文写作助手，需严格按照《党政机关公文处理工作条例》规定的通知格式撰写。内容需包含标题（发文机关+事由+文种）、主送机关、正文（明确通知缘由、具体事项、执行要求，涉及任务需标注责任主体）、附件说明（如有附件需注明附件名称及数量）、发文机关署名、成文日期、签发人（上行文需标注），语言需庄重严谨、逻辑清晰，避免歧义。",
    "请示": "你是专业的正式公文写作助手，需严格遵循请示“一文一事、单头主送”的核心要求撰写。内容需包含标题（发文机关+事由+文种）、主送机关（仅限一个直接上级机关）、正文（开门见山说明请示缘由，需引用政策依据或实际需求；清晰列出请示事项，做到具体明确；提出合理可行的建议方案供上级决策）、发文机关署名、成文日期、签发人，语言需谦恭得体，避免多头请示或夹带报告事项。",
    "会议纪要": "你是专业的正式公文写作助手，需按规范会议纪要格式撰写，确保内容客观准确、要素完整。需包含标题（会议名称+文种，如“XX会议纪要”）、基本信息（会议时间、会议地点、主持人、记录人、参会人员/列席人员/缺席人员名单及职务）、正文（分点梳理会议讨论要点，需准确反映不同观点；明确会议形成的决议事项，标注每项决议的责任单位、完成时限）、结尾（如需补充说明未尽事宜或后续沟通方式，可简要标注），语言需简洁凝练，避免主观评价。",
    "报告": "你是专业的正式公文写作助手，需严格按照报告的法定格式与写作要求撰写，区分“汇报工作”“反映情况”“答复询问”等不同报告类型。内容需包含标题（发文机关+事由+文种，如“XX关于XX工作的报告”）、主送机关、正文（开头简要说明报告背景或依据，主体部分详细阐述工作进展、成果、存在问题、原因分析及下一步计划，答复类报告需直接回应上级询问事项）、发文机关署名、成文日期、签发人（上行文需标注），语言需客观真实，不得夹带请示事项。",
    "函": "你是专业的正式公文写作助手，需根据函的“不相隶属机关商洽工作、询问答复问题、请求批准事项”功能，按规范格式撰写。内容需包含标题（发文机关+事由+文种，如“XX关于XX事宜的函”或“XX关于XX问题的复函”）、主送机关（明确的不相隶属机关）、正文（商洽类需说明事由及具体需求，询问类需列出清晰的问题清单，答复类需直接回应来函事项，请求批准类需说明依据及具体请求）、发文机关署名、成文日期，语言需平和得体，避免命令性或指令性表述。",
    "对话": "你是专业的公文场景对话撰写助手，需围绕公文办理全流程（如公文起草沟通、审批意见反馈、事项协调对接等）撰写对话内容。需明确对话场景（如“公文起草小组沟通会议对话”“上下级机关审批意见反馈对话”）、对话主体（标注角色及职务，如“起草人-XX部科员”“审批人-XX局副局长”）、对话逻辑（需符合公文办理规范，内容需聚焦具体事项，如格式修改建议、内容补充要求、办理时限确认等），语言需贴合职场沟通语境，既保持正式性，又体现沟通的针对性与高效性。"
}

# ----------------- 数据库 Session -----------------
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# ----------------- 上传模板 -----------------
@router.post("/upload-template")
async def upload_template(
    file: UploadFile = File(...),
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    path = await save_uploaded_file(file, UPLOAD_DIR)
    
    # 保存模板记录到数据库
    template = Template(
        user_id=current_user.id,
        filename=os.path.basename(path),
        original_name=file.filename
    )
    db.add(template)
    db.commit()
    db.refresh(template)
    
    return {"filename": template.filename, "original_name": template.original_name}

@router.post("/generate")
async def generate_document(
    doc_type: str = Form(...),
    user_input: str = Form(...),
    conv_id: int = Form(...),  
    template_content: str = Form(None),
    template_filename: str | None = Form(None),
    current_user = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    base_prompt = PROMPTS.get(doc_type, "请写一份正式公文：{type}")
    prompt = base_prompt.format(type=doc_type) + "\n用户要求：" + user_input

    generated = generate_text(prompt)

    # 生成 DOCX
    doc = Document()
    # 设置全局字体（仿宋_GB2312）
    style = doc.styles['Normal']
    font = style.font
    font.name = 'FangSong_GB2312'  # 仿宋是公文标准
    font.size = Pt(16)  # 小四号字
    font.bold = False
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'FangSong_GB2312')  # 中文字体
    
    # 设置段落间距
    style.paragraph_format.space_before = Pt(0)  # 段前距为 0
    style.paragraph_format.space_after = Pt(0)   # 段后距为 0

    for line in generated.split("\n"):
        line = line.strip()
        if not line:
            continue  # 跳过空行
        
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        else:
            p = doc.add_paragraph()
            parts = line.split("**")
            for i, part in enumerate(parts):
                if i % 2 == 1:
                    run = p.add_run(part)
                    run.bold = True
                else:
                    subparts = part.split("*")
                    for j, sp in enumerate(subparts):
                        run = p.add_run(sp)
                        if j % 2 == 1:
                            run.italic = True

    filename = f"{doc_type}_{int(time.time())}.docx"
    file_path = os.path.join(DOWNLOAD_DIR, filename)
    doc.save(file_path)

    # 保存到数据库
    doc_record = DocumentHistory(
        user_id=current_user.id,
        doc_type=doc_type,
        content=generated,
        filename=filename,
        template_id=template.id if template else None  # 如果有模板，则保存模板ID
    )

    db.add(doc_record)
    db.commit()
 # ✅ ========== 开始：保存对话 ==========
    try:
        data = load_data()

        # 添加用户消息
        _add_message_to_conversation(data, conv_id, "user", user_input)

        # 添加 AI 消息（带 docx 文件名）
        _add_message_to_conversation(data, conv_id, "assistant", generated, docx_file=filename)

        # 保存到文件
        save_data(data)
    except Exception as e:
        print(f"保存对话失败: {e}")
        # 不阻断主流程，继续返回
    # ✅ ========== 结束：保存对话 ==========

    return JSONResponse({
        "text": generated, 
        "filename": filename
    }, status_code=200)

# ----------------- 下载 DOCX -----------------
@router.get("/download/{filename}")
async def download(filename: str, current_user = Depends(get_current_user)):
    file_path = os.path.join(DOWNLOAD_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件未找到")
    
    # 检查文件是否属于当前用户
    doc_record = db.query(DocumentHistory).filter(DocumentHistory.filename == filename).first()
    if doc_record and doc_record.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="无权访问此文件")
    
    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )


# ----------------- 获取用户文档历史 -----------------
@router.get("/history")
async def get_history(current_user = Depends(get_current_user), db: Session = Depends(get_db)):
    docs = db.query(DocumentHistory).filter(DocumentHistory.user_id == current_user.id).order_by(DocumentHistory.created_at.desc()).all()
    return [{"id": d.id, "doc_type": d.doc_type, "filename": d.filename, "created_at": d.created_at} for d in docs]

# ----------------- 获取用户模板列表 -----------------
@router.get("/templates")
async def get_templates(current_user = Depends(get_current_user), db: Session = Depends(get_db)):
    templates = db.query(Template).filter(Template.user_id == current_user.id).order_by(Template.uploaded_at.desc()).all()
    return [{"id": t.id, "filename": t.filename, "original_name": t.original_name, "uploaded_at": t.uploaded_at} for t in templates]


@router.get("/documents")
async def list_documents():
    files = os.listdir(DOWNLOAD_DIR)
    files = [f for f in files if f.endswith(".docx")]
    return {"documents": files}

@router.get("/template-content/{filename}")
async def get_template_content(filename: str, current_user=Depends(get_current_user)):
    file_path = os.path.join(UPLOAD_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="模板文件未找到")

    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return JSONResponse({"content": "\n".join(full_text)})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"读取模板失败: {str(e)}")
    
@router.get("/keys", response_model=List[AIModelResponse])
def get_user_api_keys(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    获取当前登录用户的API Key信息（自动脱敏处理）
    """
    # 查询当前用户的所有API Key记录
    user_api_keys = db.query(AIModel).filter(
        AIModel.user_id == current_user.id  # 关联用户ID
    ).all()

    if not user_api_keys:
        return []  # 无记录时返回空列表

    # 对返回结果进行处理（如果模型中没有自动脱敏）
    for item in user_api_keys:
            # 1. 先对加密的 Key 解密，再进行脱敏处理
        if item.api_key:
            decrypted_key = decrypt_api_key(item.api_key)  # 解密加密后的 Key
            # 基于原始 Key 做脱敏（保留前4后4，中间用*代替）
            item.api_key_mask = f"{decrypted_key[:4]}****{decrypted_key[-4:]}"

    return user_api_keys

@router.get("/platforms", response_model=List[PlatformModelResponse])
async def get_platforms():
    platforms = [
        {"platform": "OpenAI", "models": ["gpt-3.5", "gpt-4"]},
        {"platform": "Anthropic", "models": ["claude-1", "claude-2"]},
        # 其他平台
    ]
    return platforms
@router.post("/keys", response_model=AIModelResponse)
async def create_ai_model(
    ai_model_create: AIModelCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    # 确保平台名称、模型名和API Key都存在
    if not ai_model_create.platform or not ai_model_create.model_name or not ai_model_create.api_key:
        raise HTTPException(status_code=422, detail="平台名称、模型名或API key不能为空")

    # 创建 AI 模型实例并存储到数据库
    # 调用 encrypt_api_key 加密原始 API Key 后存储
    encrypted_key = encrypt_api_key(ai_model_create.api_key)
    new_ai_model = AIModel(
        platform_name=ai_model_create.platform,
        model_name=ai_model_create.model_name,
        api_key=encrypted_key,  # 存储加密后的 Key
        user_id=current_user.id,
        created_at=datetime.now(pytz.UTC)
    )

    try:
        db.add(new_ai_model)
        db.commit()
        db.refresh(new_ai_model)
    except Exception as e:
        db.rollback()  # 回滚数据库操作
        raise HTTPException(status_code=500, detail="创建 AI 模型失败")

    return AIModelResponse(
        id=new_ai_model.id,
        platform=new_ai_model.platform_name,
        model_name=new_ai_model.model_name,
        api_key_mask=new_ai_model.api_key[:5] + "****",
        created_at=new_ai_model.created_at
    )
@router.put("/keys/{key_id}", response_model=AIModelResponse)
async def update_ai_model(
    key_id: int,
    ai_model_create: AIModelCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    # 查找指定 ID 的 AI 模型
    ai_model = db.query(AIModel).filter(AIModel.id == key_id, AIModel.user_id == current_user.id).first()
    
    if not ai_model:
        raise HTTPException(status_code=404, detail="API 配置未找到")

    # 确保平台名称、模型名和API Key都存在
    if not ai_model_create.platform or not ai_model_create.model_name or not ai_model_create.api_key:
        raise HTTPException(status_code=422, detail="平台名称、模型名或API key不能为空")

    # 更新 AI 模型实例
    ai_model.platform_name = ai_model_create.platform
    ai_model.model_name = ai_model_create.model_name
    # 调用 encrypt_api_key 加密新的原始 Key 后更新
    ai_model.api_key = encrypt_api_key(ai_model_create.api_key)
    ai_model.updated_at = datetime.now(pytz.UTC)

    try:
        db.commit()
        db.refresh(ai_model)
    except Exception as e:
        db.rollback()  # 回滚数据库操作
        raise HTTPException(status_code=500, detail="更新 AI 模型失败")

    return AIModelResponse(
        id=ai_model.id,
        platform=ai_model.platform_name,
        model_name=ai_model.model_name,
        api_key_mask=ai_model.api_key[:5] + "****",
        created_at=ai_model.created_at
    )
@router.delete("/keys/{key_id}", response_model=AIModelResponse)
async def delete_ai_model(
    key_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    # 查找指定 ID 的 AI 模型
    ai_model = db.query(AIModel).filter(AIModel.id == key_id, AIModel.user_id == current_user.id).first()

    if not ai_model:
        raise HTTPException(status_code=404, detail="API 配置未找到")

    try:
        db.delete(ai_model)
        db.commit()
    except Exception as e:
        db.rollback()  # 回滚数据库操作
        raise HTTPException(status_code=500, detail="删除 AI 模型失败")

    return AIModelResponse(
        id=ai_model.id,
        platform=ai_model.platform_name,
        model_name=ai_model.model_name,
        api_key_mask=ai_model.api_key[:5] + "****",
        created_at=ai_model.created_at
    )

@router.post("/delete-conversation")
async def delete_conversation(request: Request):
    data = await request.json()
    conv_id = data.get("id")

    if not conv_id:
        raise HTTPException(status_code=400, detail="缺少会话ID")

    try:
        if os.path.exists(CONVERSATIONS_FILE):
            with open(CONVERSATIONS_FILE, "r", encoding="utf-8") as f:
                conversations = json.load(f)

            new_conversations = [c for c in conversations if c["id"] != conv_id]

            if len(conversations) == len(new_conversations):
                raise HTTPException(status_code=404, detail="会话不存在")

            with open(CONVERSATIONS_FILE, "w", encoding="utf-8") as f:
                json.dump(new_conversations, f, ensure_ascii=False, indent=2)

            return JSONResponse({"message": "删除成功"})
        else:
            raise HTTPException(status_code=404, detail="会话文件不存在")

    except Exception as e:
        print(f"删除会话失败: {e}")
        raise HTTPException(status_code=500, detail="删除失败")
@router.get("/test-write")
async def test_write():
    try:
        test_file = os.path.join(DOWNLOAD_DIR, "test_write.txt")
        with open(test_file, "w", encoding="utf-8") as f:
            f.write("测试写入权限\n")
        return {"status": "success", "message": "可以写入 generated_docs"}
    except Exception as e:
        return {"status": "error", "message": str(e)}

def save_conversation(conv_id: int, user_msg: str, ai_msg: dict):
    # 加载现有会话
    if os.path.exists(CONVERSATIONS_FILE):
        with open(CONVERSATIONS_FILE, "r", encoding="utf-8") as f:
            conversations = json.load(f)
    else:
        conversations = []

    # 查找会话
    conv = next((c for c in conversations if c["id"] == conv_id), None)

    # 如果是新会话，创建一个
    if not conv:
        # 从 user_msg 提取标题（前10个字）
        title = user_msg.strip()[:10] + "..." if len(user_msg.strip()) > 10 else user_msg.strip()
        conv = {
            "id": conv_id,
            "title": f"{title}",
            "created_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            "updated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            "messages": []
        }
        conversations.append(conv)
    else:
        conv["updated_at"] = time.strftime("%Y-%m-%d %H:%M:%S")

    # 添加用户消息
    conv["messages"].append({
        "role": "user",
        "content": user_msg,
        "created_at": time.strftime("%Y-%m-%d %H:%M:%S")
    })

    # 添加 AI 消息
    conv["messages"].append({
        "role": "assistant",
        "content": ai_msg["text"],
        "docx_file": ai_msg["filename"],
        "created_at": time.strftime("%Y-%m-%d %H:%M:%S")
    })

    # 保存回文件
    with open(CONVERSATIONS_FILE, "w", encoding="utf-8") as f:
        json.dump(conversations, f, ensure_ascii=False, indent=2)